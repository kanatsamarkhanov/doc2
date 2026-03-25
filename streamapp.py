# ═══════════════════════════════════════════════════════════════════════════
#  Smart Article — Research Writing Platform  v5.0
#  streamlit run app.py
# ═══════════════════════════════════════════════════════════════════════════
import streamlit as st
import re, json, io, hashlib, base64, smtplib
from datetime import datetime
from io import BytesIO
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import pandas as pd
    PD_OK = True
except ImportError:
    PD_OK = False

try:
    import requests
    REQ_OK = True
except ImportError:
    REQ_OK = False

st.set_page_config(page_title="Smart Article", page_icon="📝",
                   layout="wide", initial_sidebar_state="expanded")

# ─── Local fallback paths ─────────────────────────────────────────────────
USERS_FILE = Path("users.json")
LOGS_FILE  = Path("logs.json")

# ═══════════════════════════════════════════════════════════════════════════
#  EMAIL NOTIFICATIONS (SMTP via st.secrets)
# ═══════════════════════════════════════════════════════════════════════════
def send_email_notification(subject: str, body_html: str):
    """Send email to admin via Gmail SMTP. Requires SMTP_USER + SMTP_PASS in secrets."""
    try:
        smtp_user  = st.secrets.get("SMTP_USER", "")
        smtp_pass  = st.secrets.get("SMTP_PASS", "")
        admin_mail = st.secrets.get("ADMIN_EMAIL", "kanat.baurzhanuly@gmail.com")
        if not smtp_user or not smtp_pass:
            return
        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"]    = smtp_user
        msg["To"]      = admin_mail
        msg.attach(MIMEText(body_html, "html"))
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as srv:
            srv.login(smtp_user, smtp_pass)
            srv.send_message(msg)
    except Exception as e:
        print(f"[SMTP] {e}")

# ═══════════════════════════════════════════════════════════════════════════
#  TRANSLATIONS  (default = KZ)
# ═══════════════════════════════════════════════════════════════════════════
TR = {
"🇰🇿 Қазақша": dict(
    app="Smart Article", tagline="Ғылыми мақала жазу платформасы",
    sign_in="Кіру", register="Тіркелу", username="Пайдаланушы аты",
    password="Құпия сөз", confirm_pw="Растаңыз", email="E-mail",
    full_name="Толық аты", role="Рөл",
    roles=["Зерттеуші","Докторант","Профессор","Аналитик","Студент","Басқа"],
    login_btn="Кіру →", reg_btn="Аккаунт жасау →",
    login_err="Қате логин немесе құпия сөз",
    pw_mismatch="Құпия сөздер сәйкес келмейді",
    pw_short="Құпия сөз ≥ 6 таңба",
    uname_short="Логин ≥ 3 таңба",
    username_taken="Мұндай логин бар",
    email_taken="E-mail тіркелген",
    reg_ok="✅ Аккаунт жасалды! Жүйеге кіріңіз.",
    logout="Шығу",
    guest_msg="🔒 Өңдеу үшін жүйеге кіріңіз немесе тіркеліңіз",
    login_required="⚠️ Бұл функцияны пайдалану үшін жүйеге кіру қажет",
    login_link="🔑 Жүйеге кіру →",
    nav_info="Ақпарат", nav_sec="Бөлімдер", nav_media="Медиа",
    nav_refs="Әдебиет", nav_gen="Экспорт", nav_thesis="Диплом", nav_set="Параметрлер",
    nav_agent="ЖИ Агент",
    title_f="Мақала атауы", authors_f="Авторлар", affil_f="Аффилиация",
    journal_f="Журнал", keywords_f="Кілт сөздер", abstract_f="Аңдатпа",
    art_type_f="Мақала түрі",
    art_types=["Ғылыми мақала","Шолу","Мини-шолу","Хабарлама","Хат","Кейс-стади"],
    w_words="Сөздер", w_secs="Бөлімдер", w_figs="Суреттер",
    w_tbls="Кестелер", w_refs="Әдебиет", w_forms="Формулалар",
    completeness="Толықтығы",
    readiness="Дайындығы",
    missing="Жетіспейді",
    cite_styles=["APA 7th","Ванкувер","Гарвард","IEEE","ГОСТ 7.0.5-2008"],
    cite_style_lbl="Дәйексөз стилі",
    add_btn="Қосу", del_btn="Жою",
    add_fig="Сурет қосу", add_tbl="Кесте қосу", add_form="Формула қосу",
    add_ref="Дереккөз қосу",
    no_figs="Суреттер жоқ.", no_tbls="Кестелер жоқ.",
    no_forms="Формулалар жоқ.", no_refs="Дереккөздер жоқ.",
    insert_fig="↩ Сурет кірістіру", insert_tbl="↩ Кесте кірістіру",
    insert_form="↩ Формула кірістіру",
    fig_lbl="Сурет", tbl_lbl="Кесте", form_lbl="Формула",
    caption_lbl="Аңыз", upload_img="Сурет жүктеу (PNG/JPG/TIF)",
    csv_data="Деректер (CSV)", latex_lbl="LaTeX коды", desc_lbl="Сипаттама",
    objectives_tool="🎯 Мақсат пен міндеттерді тексеру",
    obj_found="✅ Анықталды:", obj_missing="⚠️ Табылмады:",
    obj_tip="Мақсат пен міндеттер кіріспеде нақты тұжырымдалуы керек.",
    aim_kw=["мақсат","міндет","зерттеу мақсаты","жұмыстың мақсаты"],
    task_kw=["міндеттер","зерттеу міндеттері","жұмыс міндеттері"],
    dl_docx="📥 DOCX жүктеу", dl_md="📥 Markdown жүктеу",
    save_json="💾 Жобаны сақтау",
    load_json="📂 Жобаны жүктеу",
    reset_btn="Деректерді тазалау", reset_ok="✅ Тазаланды.",
    loaded="✅ Жоба жүктелді!",
    success_gen="✅ Сәтті жасалды!",
    warn_title="⚠️ Мақала атауын енгізіңіз.",
    theme_lbl="Тақырып",
    dark_mode="Күңгірт 🌙", light_mode="Жарық ☀️",
    word_count="Сөз саны",
    preview_lbl="Алдын ала қарау",
    sec_editor="Бөлім редакторы",
    feedback_lbl="💬 Пікір",
    feedback_ph="Платформаны жақсарту туралы ұсыныстар…",
    feedback_send="📨 Жіберу",
    feedback_ok="✅ Рахмет!",
    # Thesis
    thesis_title="Дипломдық жұмыс",
    thesis_nav="🎓 Диплом",
    th_titlepage="Титульный бет", th_abstract="Аңдатпа (3 тілде)",
    th_contents="Мазмұны", th_abbrev="Қысқартулар",
    th_intro="Кіріспе", th_ch1="1-тарау: Әдеби шолу",
    th_ch2="2-тарау: Методология", th_ch3="3-тарау: Нәтижелер",
    th_concl="Қорытынды", th_refs="Әдебиеттер тізімі",
    th_appendix="Қосымшалар",
    th_norm_check="📏 Нормативтік бақылау",
    th_pages="Беттер саны", th_font="Шрифт", th_spacing="Аралық",
    th_margin="Жиектер",
    gh_title="GitHub синхрондау",
    gh_token="GitHub Token", gh_repo="Репозиторий (owner/repo)",
    gh_save="☁️ Сақтау", gh_load="📥 Жүктеу",
    gh_url="Gist URL", gh_saved="✅ Сақталды!", gh_loaded="✅ Жүктелді!",
    gh_err="GitHub қатесі",
    import_bibtex="BibTeX импорты", import_plain="Мәтін ретінде",
    imported="Импортталды", ref_s="дереккөз",
    ref_type="Түрі", ref_au="Авторлар", ref_yr="Жыл",
    ref_ti="Атауы", ref_jn="Журнал / Баспа",
    ref_vol="Том", ref_no="Нөмір", ref_pp="Беттер", ref_doi="DOI",
    ref_city="Қала", ref_pub="Баспа",
    ref_types=["Журнал мақаласы","Кітап","Кітап тарауы",
               "Конференция баяндамасы","Сайт","Диссертация"],
    drag_drop="Файлды осында сүйреп апарыңыз",
    file_limit="Файл шектеуі: 200МБ",
),
"🇷🇺 Русский": dict(
    app="Smart Article", tagline="Платформа написания научных статей",
    sign_in="Вход", register="Регистрация", username="Логин",
    password="Пароль", confirm_pw="Подтвердите", email="E-mail",
    full_name="Полное имя", role="Роль",
    roles=["Исследователь","Докторант","Профессор","Аналитик","Студент","Другое"],
    login_btn="Войти →", reg_btn="Создать аккаунт →",
    login_err="Неверный логин или пароль",
    pw_mismatch="Пароли не совпадают",
    pw_short="Пароль ≥ 6 символов",
    uname_short="Логин ≥ 3 символов",
    username_taken="Логин уже занят",
    email_taken="E-mail уже зарегистрирован",
    reg_ok="✅ Аккаунт создан! Войдите в систему.",
    logout="Выйти",
    guest_msg="🔒 Войдите или зарегистрируйтесь для редактирования",
    login_required="⚠️ Для использования этой функции необходимо войти в систему",
    login_link="🔑 Войти в систему →",
    nav_info="Информация", nav_sec="Разделы", nav_media="Медиа",
    nav_refs="Литература", nav_gen="Экспорт", nav_thesis="Диплом", nav_set="Настройки",
    nav_agent="ИИ Агент",
    title_f="Название статьи", authors_f="Авторы", affil_f="Аффилиация",
    journal_f="Журнал", keywords_f="Ключевые слова", abstract_f="Аннотация",
    art_type_f="Тип статьи",
    art_types=["Научная статья","Обзор","Мини-обзор","Сообщение","Письмо","Кейс-стади"],
    w_words="Слова", w_secs="Разделы", w_figs="Рисунки",
    w_tbls="Таблицы", w_refs="Источники", w_forms="Формулы",
    completeness="Заполненность",
    readiness="Готовность",
    missing="Отсутствует",
    cite_styles=["APA 7th","Ванкувер","Гарвард","IEEE","ГОСТ 7.0.5-2008"],
    cite_style_lbl="Стиль цитирования",
    add_btn="Добавить", del_btn="Удалить",
    add_fig="Добавить рисунок", add_tbl="Добавить таблицу", add_form="Добавить формулу",
    add_ref="Добавить источник",
    no_figs="Рисунков нет.", no_tbls="Таблиц нет.",
    no_forms="Формул нет.", no_refs="Источников нет.",
    insert_fig="↩ Вставить рисунок", insert_tbl="↩ Вставить таблицу",
    insert_form="↩ Вставить формулу",
    fig_lbl="Рисунок", tbl_lbl="Таблица", form_lbl="Формула",
    caption_lbl="Подпись", upload_img="Загрузить изображение",
    csv_data="Данные (CSV)", latex_lbl="Код LaTeX", desc_lbl="Описание",
    objectives_tool="🎯 Проверка цели и задач",
    obj_found="✅ Найдено:", obj_missing="⚠️ Не найдено:",
    obj_tip="Цель и задачи должны быть чётко сформулированы во введении.",
    aim_kw=["цель","задачи","цель работы","цель исследования"],
    task_kw=["задачи","задачи исследования","задачи работы"],
    dl_docx="📥 Скачать DOCX", dl_md="📥 Скачать Markdown",
    save_json="💾 Сохранить проект",
    load_json="📂 Загрузить проект",
    reset_btn="Очистить данные", reset_ok="✅ Данные очищены.",
    loaded="✅ Проект загружен!",
    success_gen="✅ Сгенерировано!",
    warn_title="⚠️ Введите название статьи.",
    theme_lbl="Тема",
    dark_mode="Тёмная 🌙", light_mode="Светлая ☀️",
    word_count="Слов",
    preview_lbl="Предпросмотр",
    sec_editor="Редактор раздела",
    feedback_lbl="💬 Обратная связь",
    feedback_ph="Предложения по улучшению платформы…",
    feedback_send="📨 Отправить",
    feedback_ok="✅ Спасибо!",
    thesis_title="Дипломная работа",
    thesis_nav="🎓 Диплом",
    th_titlepage="Титульный лист", th_abstract="Реферат (3 языка)",
    th_contents="Содержание", th_abbrev="Обозначения и сокращения",
    th_intro="Введение", th_ch1="Глава 1: Обзор литературы",
    th_ch2="Глава 2: Методология", th_ch3="Глава 3: Результаты",
    th_concl="Заключение", th_refs="Список источников",
    th_appendix="Приложения",
    th_norm_check="📏 Нормоконтроль",
    th_pages="Страниц", th_font="Шрифт", th_spacing="Интервал",
    th_margin="Поля",
    gh_title="Синхронизация GitHub",
    gh_token="GitHub Token", gh_repo="Репозиторий (owner/repo)",
    gh_save="☁️ Сохранить", gh_load="📥 Загрузить",
    gh_url="Gist URL", gh_saved="✅ Сохранено!", gh_loaded="✅ Загружено!",
    gh_err="Ошибка GitHub",
    import_bibtex="Импорт BibTeX", import_plain="Обычный текст",
    imported="Импортировано", ref_s="источник(ов)",
    ref_type="Тип", ref_au="Авторы", ref_yr="Год",
    ref_ti="Название", ref_jn="Журнал / Издательство",
    ref_vol="Том", ref_no="Номер", ref_pp="Страницы", ref_doi="DOI",
    ref_city="Город", ref_pub="Издательство",
    ref_types=["Журнальная статья","Книга","Глава книги",
               "Материалы конференции","Сайт","Диссертация"],
    drag_drop="Перетащите файл сюда",
    file_limit="Ограничение: 200МБ",
),
"🇬🇧 English": dict(
    app="Smart Article", tagline="Research Writing Platform",
    sign_in="Sign In", register="Register", username="Username",
    password="Password", confirm_pw="Confirm", email="E-mail",
    full_name="Full Name", role="Role",
    roles=["Researcher","PhD Student","Professor","Analyst","Student","Other"],
    login_btn="Sign In →", reg_btn="Create Account →",
    login_err="Invalid credentials",
    pw_mismatch="Passwords do not match",
    pw_short="Password must be ≥ 6 characters",
    uname_short="Username must be ≥ 3 characters",
    username_taken="Username already taken",
    email_taken="E-mail already registered",
    reg_ok="✅ Account created! You can sign in now.",
    logout="Logout",
    guest_msg="🔒 Sign in or register to start editing",
    login_required="⚠️ Please sign in to use this feature",
    login_link="🔑 Sign In →",
    nav_info="Info", nav_sec="Sections", nav_media="Media",
    nav_refs="Refs", nav_gen="Export", nav_thesis="Thesis", nav_set="Settings",
    nav_agent="AI Agent",
    title_f="Article Title", authors_f="Authors", affil_f="Affiliation",
    journal_f="Journal", keywords_f="Keywords", abstract_f="Abstract",
    art_type_f="Article Type",
    art_types=["Research Article","Review","Mini-review","Communication","Letter","Case Study"],
    w_words="Words", w_secs="Sections", w_figs="Figures",
    w_tbls="Tables", w_refs="Refs", w_forms="Formulas",
    completeness="Completeness",
    readiness="Readiness",
    missing="Missing",
    cite_styles=["APA 7th","Vancouver","Harvard","IEEE","ГОСТ 7.0.5-2008"],
    cite_style_lbl="Citation Style",
    add_btn="Add", del_btn="Delete",
    add_fig="Add Figure", add_tbl="Add Table", add_form="Add Formula",
    add_ref="Add Reference",
    no_figs="No figures yet.", no_tbls="No tables yet.",
    no_forms="No formulas yet.", no_refs="No references yet.",
    insert_fig="↩ Insert Figure", insert_tbl="↩ Insert Table",
    insert_form="↩ Insert Formula",
    fig_lbl="Figure", tbl_lbl="Table", form_lbl="Formula",
    caption_lbl="Caption", upload_img="Upload image (PNG/JPG/TIF)",
    csv_data="Data (CSV)", latex_lbl="LaTeX code", desc_lbl="Description",
    objectives_tool="🎯 Check Objectives & Tasks",
    obj_found="✅ Found:", obj_missing="⚠️ Missing:",
    obj_tip="Objectives and tasks must be clearly stated in the Introduction.",
    aim_kw=["objective","aim","purpose","goal","research objective"],
    task_kw=["tasks","research tasks","objectives of the study"],
    dl_docx="📥 Download DOCX", dl_md="📥 Download Markdown",
    save_json="💾 Save Project",
    load_json="📂 Load Project",
    reset_btn="Reset all data", reset_ok="✅ Cleared.",
    loaded="✅ Project loaded!",
    success_gen="✅ Generated!",
    warn_title="⚠️ Please add an article title first.",
    theme_lbl="Theme",
    dark_mode="Dark 🌙", light_mode="Light ☀️",
    word_count="Words",
    preview_lbl="Preview",
    sec_editor="Section Editor",
    feedback_lbl="💬 Feedback",
    feedback_ph="Your suggestions to improve the platform…",
    feedback_send="📨 Send",
    feedback_ok="✅ Thank you!",
    thesis_title="Diploma / Thesis",
    thesis_nav="🎓 Thesis",
    th_titlepage="Title Page", th_abstract="Abstract (3 languages)",
    th_contents="Table of Contents", th_abbrev="Abbreviations",
    th_intro="Introduction", th_ch1="Chapter 1: Literature Review",
    th_ch2="Chapter 2: Methodology", th_ch3="Chapter 3: Results",
    th_concl="Conclusion", th_refs="References",
    th_appendix="Appendices",
    th_norm_check="📏 Normative Check",
    th_pages="Pages", th_font="Font", th_spacing="Line Spacing",
    th_margin="Margins",
    gh_title="GitHub Sync",
    gh_token="GitHub Token", gh_repo="Repository (owner/repo)",
    gh_save="☁️ Save", gh_load="📥 Load",
    gh_url="Gist URL", gh_saved="✅ Saved!", gh_loaded="✅ Loaded!",
    gh_err="GitHub error",
    import_bibtex="Import BibTeX", import_plain="Plain text",
    imported="Imported", ref_s="reference(s)",
    ref_type="Type", ref_au="Authors", ref_yr="Year",
    ref_ti="Title", ref_jn="Journal / Publisher",
    ref_vol="Volume", ref_no="Number", ref_pp="Pages", ref_doi="DOI",
    ref_city="City", ref_pub="Publisher",
    ref_types=["Journal Article","Book","Book Chapter",
               "Conference Paper","Website","Thesis"],
    drag_drop="Drag and drop file here",
    file_limit="File limit: 200MB",
),
}

# ═══════════════════════════════════════════════════════════════════════════
#  TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════
TEMPLATES = {
    "ENU Journal (KZ)": {
        "lang":"KZ","cite_style":"APA 7th","abstract_max":300,"keywords_count":"3–10",
        "irsti_label":"ХҒТАР","irsti_hint":"grnti.ru арқылы ХҒТАР кодын көрсетіңіз",
        "sections":["intro","materials_methods","results","discussion","conclusion",
                    "supplementary","author_contributions","author_info","funding","acknowledgements","conflicts"],
    },
    "ENU Journal (RU)": {
        "lang":"RU","cite_style":"APA 7th","abstract_max":300,"keywords_count":"3–10",
        "irsti_label":"МРНТИ","irsti_hint":"Укажите код МРНТИ через grnti.ru",
        "sections":["intro","materials_methods","results","discussion","conclusion",
                    "supplementary","author_contributions","author_info","funding","acknowledgements","conflicts"],
    },
    "ENU Journal (EN)": {
        "lang":"EN","cite_style":"APA 7th","abstract_max":300,"keywords_count":"3–10",
        "irsti_label":"IRSTI","irsti_hint":"Specify IRSTI code via grnti.ru",
        "sections":["intro","materials_methods","results","discussion","conclusion",
                    "supplementary","author_contributions","author_info","funding","acknowledgements","conflicts"],
        "trilingual":True,
    },
    "Generic Research Article": {
        "lang":"EN","cite_style":"APA 7th","abstract_max":250,"keywords_count":"5–8",
        "irsti_label":"","irsti_hint":"",
        "sections":["intro","materials_methods","results","discussion","conclusion","acknowledgements","conflicts"],
    },
    "IEEE Conference Paper": {
        "lang":"EN","cite_style":"IEEE","abstract_max":150,"keywords_count":"4–6",
        "irsti_label":"","irsti_hint":"",
        "sections":["intro","materials_methods","results","discussion","conclusion","acknowledgements"],
    },
}

ALL_SECTIONS = {
    "intro":               {"icon":"1️⃣","kz":"Кіріспе","ru":"Введение","en":"Introduction"},
    "materials_methods":   {"icon":"2️⃣","kz":"Материалдар мен әдістер","ru":"Материалы и методы","en":"Materials & Methods"},
    "results":             {"icon":"3️⃣","kz":"Нәтижелер","ru":"Результаты","en":"Results"},
    "discussion":          {"icon":"4️⃣","kz":"Талқылау","ru":"Обсуждение","en":"Discussion"},
    "conclusion":          {"icon":"5️⃣","kz":"Қорытынды","ru":"Заключение","en":"Conclusion"},
    "supplementary":       {"icon":"📎","kz":"Қосымша материалдар","ru":"Дополнительные материалы","en":"Supplementary Materials"},
    "author_contributions":{"icon":"👥","kz":"Авторлар үлесі","ru":"Вклад авторов","en":"Author Contributions"},
    "author_info":         {"icon":"👤","kz":"Авторлар туралы","ru":"Об авторах","en":"Author Information"},
    "funding":             {"icon":"💰","kz":"Қаржыландыру","ru":"Финансирование","en":"Funding"},
    "acknowledgements":    {"icon":"🙏","kz":"Алғыс","ru":"Благодарности","en":"Acknowledgements"},
    "conflicts":           {"icon":"⚖️","kz":"Мүдделер қақтығысы","ru":"Конфликт интересов","en":"Conflicts of Interest"},
}

THESIS_SECTIONS = {
    "th_titlepage":   {"icon":"📋","kz":"Титульный бет","ru":"Титульный лист","en":"Title Page"},
    "th_declaration": {"icon":"✍️","kz":"Мәлімдеме","ru":"Декларация","en":"Declaration"},
    "th_abstract_kz": {"icon":"🇰🇿","kz":"Аңдатпа (ҚЗ)","ru":"Аннотация (КЗ)","en":"Abstract (KZ)"},
    "th_abstract_ru": {"icon":"🇷🇺","kz":"Аңдатпа (РУ)","ru":"Аннотация (РУ)","en":"Abstract (RU)"},
    "th_abstract_en": {"icon":"🇬🇧","kz":"Аңдатпа (АҒ)","ru":"Аннотация (АН)","en":"Abstract (EN)"},
    "th_abbrev":      {"icon":"📖","kz":"Қысқартулар","ru":"Сокращения","en":"Abbreviations"},
    "th_intro":       {"icon":"📌","kz":"Кіріспе","ru":"Введение","en":"Introduction"},
    "th_ch1":         {"icon":"📚","kz":"1-тарау","ru":"Глава 1","en":"Chapter 1"},
    "th_ch2":         {"icon":"🔬","kz":"2-тарау","ru":"Глава 2","en":"Chapter 2"},
    "th_ch3":         {"icon":"📊","kz":"3-тарау","ru":"Глава 3","en":"Chapter 3"},
    "th_ch4":         {"icon":"💡","kz":"4-тарау (міндетті емес)","ru":"Глава 4 (опц.)","en":"Chapter 4 (opt.)"},
    "th_concl":       {"icon":"🏁","kz":"Қорытынды","ru":"Заключение","en":"Conclusion"},
    "th_appendix_a":  {"icon":"📎","kz":"А Қосымша","ru":"Приложение А","en":"Appendix A"},
    "th_appendix_b":  {"icon":"📎","kz":"Б Қосымша","ru":"Приложение Б","en":"Appendix B"},
}

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION HINTS  (AI Q&A content)
# ═══════════════════════════════════════════════════════════════════════════
SEC_HINTS = {
    "intro": {
        "tip": {"kz":"Кіріспе зерттеуді кеңірек контекстке орналастырып, маңыздылығын атап өтуі керек. Зерттеу мақсаты мен міндеттерін нақты тұжырымдаңыз.",
                "ru":"Введение должно помещать исследование в контекст и подчёркивать его значимость. Чётко сформулируйте цель и задачи.",
                "en":"Introduction should situate the study in context and emphasise its significance. Clearly state the objective and tasks."},
        "qa": [
            {"q":{"kz":"Кіріспенің оңтайлы ұзындығы қандай?","ru":"Какой оптимальный объём введения?","en":"What is the optimal length of Introduction?"},
             "a":{"kz":"Негізінен 300–600 сөз. CARS моделін пайдаланыңыз: аумақты белгілеу → нишаны анықтау → нишаны иелену.","ru":"Обычно 300–600 слов. Используйте модель CARS: занять территорию → создать нишу → заполнить нишу.","en":"Typically 300–600 words. Use the CARS model: Establish Territory → Establish Niche → Occupy Niche."}},
            {"q":{"kz":"Мақсат пен міндеттерді қалай тұжырымдауға болады?","ru":"Как правильно сформулировать цель и задачи?","en":"How to formulate objectives and tasks?"},
             "a":{"kz":"Мақсат — бір, нақты, өлшенетін. Міндеттер — мақсатқа жету жолдары (3–5 тармақ). 'Зерттеу мақсаты — ...' деп бастаңыз.","ru":"Цель — одна, конкретная, измеримая. Задачи — шаги к цели (3–5 пунктов). Начните: 'Целью работы является...'","en":"One specific, measurable objective. Tasks are steps toward it (3–5 items). Start: 'The objective of this study is to...'"}},
        ]
    },
    "materials_methods": {
        "tip": {"kz":"Жұмысты қайталауға жеткілікті мәліметтер беріңіз. Бағдарламалық жасақтаманың нұсқасын, деректер көзін, статистикалық әдістерді атаңыз.",
                "ru":"Предоставьте достаточно деталей для воспроизведения работы. Укажите версию ПО, источники данных, статистические методы.",
                "en":"Provide enough detail for reproducibility. State software version, data sources, and statistical methods."},
        "qa": [
            {"q":{"kz":"ЖИ құралдары туралы не айту керек?","ru":"Что нужно написать об использовании ИИ-инструментов?","en":"What to write about AI tool usage?"},
             "a":{"kz":"Қандай ЖИ құралы (мысалы, ChatGPT), қандай мақсатта (мәтін жазу, аудару, деректерді талдау) пайдаланылғанын нақты атаңыз.","ru":"Укажите конкретный ИИ-инструмент (например, ChatGPT 4.0), цель использования (написание текста, перевод, анализ данных).","en":"Specify which AI tool (e.g., ChatGPT 4.0), for what purpose (text writing, translation, data analysis)."}},
            {"q":{"kz":"Деректер жинаудың ең жақсы тәсілі?","ru":"Как лучше описать сбор данных?","en":"How to best describe data collection?"},
             "a":{"kz":"Деректер жинау кезеңдерін: кезең → сипаттама → деректер форматы → санаттар — кезекпен сипаттаңыз.","ru":"Опишите поэтапно: этап → описание → формат данных → категории.","en":"Describe step-by-step: stage → description → data format → categories."}},
        ]
    },
    "results": {
        "tip": {"kz":"Нәтижелерді нақты жеткізіңіз. Сандық деректерді кестелер мен суреттер арқылы қолдаңыз.","ru":"Чётко изложите результаты, поддержите их таблицами и рисунками.","en":"Present results clearly, supported by tables and figures."},
        "qa": [
            {"q":{"kz":"Нәтижелерді қалай құрылымдауға болады?","ru":"Как структурировать результаты?","en":"How to structure the Results section?"},
             "a":{"kz":"Зерттеу сұрақтары немесе міндеттер бойынша топтаңыз. Әр бөлімше бір міндетті қарастырсын.","ru":"Сгруппируйте по исследовательским вопросам или задачам. Каждый подраздел — одна задача.","en":"Group by research questions or tasks. Each subsection covers one task."}},
        ]
    },
    "conclusion": {
        "tip": {"kz":"МІНДЕТТІ БӨЛІМ. Зерттеудің негізгі нәтижелерін жинақтаңыз, маңыздылығын атаңыз, болашақ зерттеулер бағытын ұсыныңыз.","ru":"ОБЯЗАТЕЛЬНЫЙ РАЗДЕЛ. Обобщите результаты, укажите значимость, предложите направления будущих исследований.","en":"MANDATORY. Summarise key findings, state significance, suggest future research directions."},
        "qa": [
            {"q":{"kz":"Қорытынды қанша сөзден тұруы керек?","ru":"Какой объём заключения?","en":"How long should the Conclusion be?"},
             "a":{"kz":"200–400 сөз. 4 бөлік: 1) не жасалды, 2) негізгі нәтижелер, 3) маңыздылық, 4) болашақ зерттеулер.","ru":"200–400 слов. 4 части: 1) что было сделано, 2) ключевые результаты, 3) значимость, 4) будущие исследования.","en":"200–400 words. 4 parts: 1) what was done, 2) key findings, 3) significance, 4) future research."}},
        ]
    },
    "author_contributions": {
        "tip": {"kz":"CRediT таксономиясы бойынша әр автордың үлесін сипаттаңыз.","ru":"Опишите вклад каждого автора по таксономии CRediT.","en":"Describe each author's contribution using CRediT taxonomy."},
        "qa": [
            {"q":{"kz":"CRediT рөлдері қандай?","ru":"Какие роли в CRediT?","en":"What are CRediT roles?"},
             "a":{"kz":"Концептуализация · Методология · Бағдарламалық жасақтама · Тексеру · Формальды талдау · Зерттеу · Ресурстар · Деректерді курация · Бастапқы мәтін жазу · Редакциялау · Визуализация · Ғылыми басшылық · Жоба басқарылуы · Қаржыландыру алу","ru":"Концептуализация · Методология · Программное обеспечение · Валидация · Формальный анализ · Исследование · Ресурсы · Курирование данных · Написание рукописи · Рецензирование · Визуализация · Руководство · Управление проектом · Получение финансирования","en":"Conceptualization · Methodology · Software · Validation · Formal analysis · Investigation · Resources · Data Curation · Writing – Original Draft · Writing – Review & Editing · Visualization · Supervision · Project administration · Funding acquisition"}},
        ]
    },
    "funding": {
        "tip": {"kz":"Сыртқы қаржыландыру туралы ақпарат беріңіз немесе 'Зерттеу сыртқы қаржыландыруды алмады' деп жазыңыз.","ru":"Укажите источники финансирования или напишите 'Исследование не получало внешнего финансирования'.","en":"State funding sources or write 'This research received no external funding'."},
        "qa": [
            {"q":{"kz":"Гранттың нөмірін қалай форматтауға болады?","ru":"Как форматировать номер гранта?","en":"How to format grant number?"},
             "a":{"kz":"'Бұл зерттеу [ҰЙЫМ] қаржыландырды, грант №AP#######.' деп жазыңыз.","ru":"'Исследование финансировалось [ОРГАНИЗАЦИЯ], грант №AP#######.'","en":"'This research was funded by [FUNDER], grant number AP#######.'"}},
        ]
    },
}

# ═══════════════════════════════════════════════════════════════════════════
#  HELPERS
# ═══════════════════════════════════════════════════════════════════════════
def t(key: str) -> str:
    lang = st.session_state.get("lang", "🇰🇿 Қазақша")
    return TR.get(lang, TR["🇰🇿 Қазақша"]).get(key, key)

def wc(text: str) -> int:
    return len(re.findall(r"\w+", text)) if text else 0

def sfn(title: str) -> str:
    return re.sub(r"[^\w\s-]", "", title)[:50].strip().replace(" ", "_") or "article"

def sec_name(key: str) -> str:
    s    = ALL_SECTIONS.get(key) or THESIS_SECTIONS.get(key, {})
    lang = st.session_state.get("lang", "🇰🇿 Қазақша")
    if "Русский" in lang: return s.get("ru", key)
    if "Қазақша" in lang: return s.get("kz", key)
    return s.get("en", key)

def is_logged_in() -> bool:
    return bool(st.session_state.get("logged_in"))

def _colors() -> dict:
    """Return theme-aware color dict for use in HTML strings."""
    dark = st.session_state.get("dark", True)
    return {
        "fg":    "#f1f5f9" if dark else "#0d1f3c",
        "sub":   "#94a3b8" if dark else "#475569",
        "muted": "#64748b" if dark else "#6b7280",
        "acc":   "#3b82f6" if dark else "#1d4ed8",
        "mbg":   "#1a2a42" if dark else "#e8eef7",
        "brd":   "#334155" if dark else "#cbd5e1",
        "cbg":   "#1e293b" if dark else "#ffffff",
    }

def current_template() -> dict:
    return TEMPLATES.get(st.session_state.get("template","ENU Journal (KZ)"),
                         TEMPLATES["ENU Journal (KZ)"])

def active_sections() -> list:
    return current_template()["sections"]

def _lang_key() -> str:
    lang = st.session_state.get("lang","🇰🇿 Қазақша")
    if "Русский" in lang: return "ru"
    if "Қазақша" in lang: return "kz"
    return "en"

# ═══════════════════════════════════════════════════════════════════════════
#  GITHUB SYNC  (users.json / logs.json → repo)
# ═══════════════════════════════════════════════════════════════════════════
def _gh_headers() -> dict:
    tok = st.session_state.get("gh_token","")
    try: tok = tok or st.secrets.get("GH_TOKEN","")
    except: pass
    return {"Authorization": f"token {tok}", "Accept": "application/vnd.github+json"} if tok else {}

def _gh_repo() -> str:
    r = st.session_state.get("gh_repo","")
    try: r = r or st.secrets.get("GH_REPO","")
    except: pass
    return r

def gh_read_file(path: str) -> tuple:
    """Returns (content_str, sha) or (None, None)."""
    if not REQ_OK: return None, None
    hdrs = _gh_headers(); repo = _gh_repo()
    if not hdrs or not repo: return None, None
    try:
        r = requests.get(f"https://api.github.com/repos/{repo}/contents/{path}",
                         headers=hdrs, timeout=8)
        if r.status_code == 200:
            data = r.json()
            content = base64.b64decode(data["content"]).decode("utf-8")
            return content, data.get("sha")
    except Exception:
        pass
    return None, None

def gh_write_file(path: str, content: str, msg: str) -> bool:
    if not REQ_OK: return False
    hdrs = _gh_headers(); repo = _gh_repo()
    if not hdrs or not repo: return False
    _, sha = gh_read_file(path)
    encoded = base64.b64encode(content.encode("utf-8")).decode("utf-8")
    payload = {"message": msg, "content": encoded}
    if sha: payload["sha"] = sha
    try:
        r = requests.put(f"https://api.github.com/repos/{repo}/contents/{path}",
                         headers=hdrs, json=payload, timeout=10)
        return r.status_code in (200, 201)
    except Exception:
        return False

# ═══════════════════════════════════════════════════════════════════════════
#  AUTH / LOGS (local + GitHub)
# ═══════════════════════════════════════════════════════════════════════════
def _hp(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def load_users() -> dict:
    # Try GitHub first
    content, _ = gh_read_file("data/users.json")
    if content:
        try: return json.loads(content)
        except: pass
    # Fallback local
    if USERS_FILE.exists():
        try: return json.loads(USERS_FILE.read_text("utf-8"))
        except: pass
    seed = {"admin": {"password": _hp("admin123"), "name": "Admin",
                      "email": "admin@smart-article.kz", "role": "Researcher",
                      "created_at": datetime.now().isoformat()}}
    _save_users(seed)
    return seed

def _save_users(u: dict):
    txt = json.dumps(u, ensure_ascii=False, indent=2)
    USERS_FILE.write_text(txt, "utf-8")
    gh_write_file("data/users.json", txt,
                  f"update users {datetime.now().strftime('%Y-%m-%d %H:%M')}")

def load_logs() -> list:
    content, _ = gh_read_file("data/logs.json")
    if content:
        try: return json.loads(content)
        except: pass
    if LOGS_FILE.exists():
        try: return json.loads(LOGS_FILE.read_text("utf-8"))
        except: pass
    return []

def add_log(event: str, user: str, detail: str = ""):
    logs = load_logs()
    logs.append({"event": event, "username": user,
                 "detail": detail, "ts": datetime.now().isoformat()})
    txt = json.dumps(logs[-2000:], ensure_ascii=False, indent=2)
    LOGS_FILE.write_text(txt, "utf-8")
    gh_write_file("data/logs.json", txt,
                  f"log {event} {user} {datetime.now().strftime('%Y-%m-%d %H:%M')}")

def do_register(uname, email, pw, name, role):
    if len(uname) < 3: return False, "uname_short"
    if len(pw) < 6:    return False, "pw_short"
    users = load_users()
    if uname in users: return False, "username_taken"
    if any(v.get("email") == email for v in users.values()): return False, "email_taken"
    users[uname] = {"password": _hp(pw), "name": name, "email": email,
                    "role": role, "created_at": datetime.now().isoformat()}
    _save_users(users)
    add_log("register", uname, f"role={role}")
    send_email_notification(
        "New Registration — Smart Article",
        f"<h2>New user registered</h2>"
        f"<p><b>Username:</b> {uname}</p><p><b>Name:</b> {name}</p>"
        f"<p><b>Email:</b> {email}</p><p><b>Role:</b> {role}</p>"
        f"<p><b>Time:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>")
    return True, "ok"

def do_login(uname, pw):
    users = load_users()
    if uname not in users:
        add_log("fail", uname, "not_found"); return False, {}
    if users[uname]["password"] != _hp(pw):
        add_log("fail", uname, "wrong_pw"); return False, {}
    add_log("login", uname)
    send_email_notification(
        f"Login — Smart Article: {uname}",
        f"<h2>User logged in</h2>"
        f"<p><b>Username:</b> {uname}</p>"
        f"<p><b>Time:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>")
    return True, users[uname]

# ═══════════════════════════════════════════════════════════════════════════
#  SESSION STATE
# ═══════════════════════════════════════════════════════════════════════════
def init_state():
    d = {
        "logged_in": False, "username": "", "user_data": {},
        "lang": "🇰🇿 Қазақша", "dark": True, "page": "info",
        "template": "ENU Journal (KZ)", "cite_style": "APA 7th",
        "art_title":"","authors":"","affiliation":"","journal":"",
        "keywords":"","abstract":"","art_type":"","irsti":"","section_field":"",
        "abstract_kz":"","abstract_ru":"",
        **{k:"" for k in ALL_SECTIONS},
        **{k:"" for k in THESIS_SECTIONS},
        "th_student":"","th_supervisor":"","th_year":str(datetime.now().year),
        "th_university":"","th_faculty":"","th_specialty":"","th_degree":"",
        "figures":[],"tables":[],"formulas":[],"refs":[],
        "gh_token":"","gh_repo":"",
        "anthropic_api_key": "",
        "show_login_msg": False,
    }
    for k, v in d.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ═══════════════════════════════════════════════════════════════════════════
#  CSS  (full dark / light adaptation)
# ═══════════════════════════════════════════════════════════════════════════
def inject_css(dark: bool = True):
    if dark:
        bg="#0f172a"; sbg="#1e293b"; cbg="#1e293b"; brd="#334155"
        fg="#f1f5f9"; sub="#94a3b8"; acc="#3b82f6"; acc2="#2563eb"
        ibg="#0f172a"; mbg="#1a2a42"; met="#60a5fa"
        inp_ph="#64748b"; code_bg="#1e293b"
        nav_act_bg=acc; nav_hover="#243551"
        hint_bg="#1a2a42"; hint_brd="#2d4a6e"; hint_fg="#93c5fd"
        btn_dl_bg="#065f46"; btn_dl_hov="#047857"
        lock_bg="#1e293b"; lock_fg="#f59e0b"; lock_brd="#f59e0b"
        tab_bg="#1a2a42"; preview_bg="#1e293b"
        insert_bg="#243551"; insert_fg="#60a5fa"
        obj_ok="#1c3a2e"; obj_ok_fg="#34d399"
        obj_miss="#3a1c1c"; obj_miss_fg="#f87171"
        # browse button
        browse_bg="#1e293b"; browse_fg="#93c5fd"; browse_brd="#334155"; browse_hov="#243551"
        # sidebar section nav
        sb_nav_bg="transparent"; sb_nav_fg="#94a3b8"; sb_nav_hov="#243551"; sb_nav_act="#1d3a5e"
        # ai badge
        ai_bg="#1e1b4b"; ai_fg="#a5b4fc"; ai_brd="#4338ca"
    else:
        bg="#f0f4f8"; sbg="#ffffff"; cbg="#ffffff"; brd="#cbd5e1"
        fg="#0d1f3c"; sub="#475569"; acc="#1d4ed8"; acc2="#1e40af"
        ibg="#ffffff"; mbg="#e8eef7"; met="#1d4ed8"
        inp_ph="#6b7280"; code_bg="#f1f5f9"
        nav_act_bg=acc; nav_hover="#dde8f5"
        hint_bg="#eff6ff"; hint_brd="#bfdbfe"; hint_fg="#1e40af"
        btn_dl_bg="#047857"; btn_dl_hov="#065f46"
        lock_bg="#fffbeb"; lock_fg="#92400e"; lock_brd="#fcd34d"
        tab_bg="#e8eef7"; preview_bg="#ffffff"
        insert_bg="#dbeafe"; insert_fg="#1d4ed8"
        obj_ok="#d1fae5"; obj_ok_fg="#065f46"
        obj_miss="#fee2e2"; obj_miss_fg="#991b1b"
        # browse button
        browse_bg="#f8fafc"; browse_fg="#1d4ed8"; browse_brd="#cbd5e1"; browse_hov="#dbeafe"
        # sidebar section nav
        sb_nav_bg="transparent"; sb_nav_fg="#475569"; sb_nav_hov="#dde8f5"; sb_nav_act="#dbeafe"
        # ai badge
        ai_bg="#ede9fe"; ai_fg="#4c1d95"; ai_brd="#7c3aed"

    # Translated drag-drop strings for CSS content injection
    drag_txt   = t("drag_drop").replace("'", "\\'")
    limit_txt  = t("file_limit").replace("'", "\\'")
    # Light/dark sensitive colours for CSS custom properties
    lm_patch = "" if dark else f"""
/* ── Light mode: force colour-scheme so browser uses light rendering ── */
:root {{ color-scheme: light; }}
*:not([class*="stButton"]):not(button):not([data-testid*="baseButton"]) {{
    color: {fg};
}}
"""

    st.markdown(f"""<style>
{lm_patch}
html,body,[data-testid="stAppViewContainer"],[data-testid="stMain"],.main .block-container{{
    background-color:{bg}!important;color:{fg}!important;
    font-family:'Inter','Segoe UI',sans-serif;
}}
[data-testid="stSidebar"],[data-testid="stSidebar"]>div{{
    background-color:{sbg}!important;border-right:1px solid {brd};
}}
/* Universal text — covers both modes */
h1,h2,h3,h4,h5,h6,p,span,li,div,label,td,th,
[data-testid="stMarkdownContainer"]*,
[data-testid="stWidgetLabel"] p,[data-testid="stWidgetLabel"] span,
[data-testid="stWidgetLabel"] label,
[data-testid="stSidebar"] p,[data-testid="stSidebar"] span,
[data-testid="stSidebar"] div,[data-testid="stSidebar"] label{{
    color:{fg}!important;
}}
[data-testid="stCaptionContainer"] *,.stCaption,small{{color:{sub}!important;}}

/* ── Inputs ────────────────────────────────────────────────────── */
.stTextInput input,.stTextArea textarea,
[data-baseweb="input"] input,[data-baseweb="textarea"] textarea,
[data-baseweb="base-input"] input{{
    background-color:{ibg}!important;color:{fg}!important;
    border:1.5px solid {brd}!important;border-radius:8px!important;
    caret-color:{fg}!important;font-size:1rem!important;
}}
.stTextInput input::placeholder,.stTextArea textarea::placeholder{{color:{inp_ph}!important;opacity:1!important;}}
.stTextInput input:focus,.stTextArea textarea:focus{{
    border-color:{acc}!important;box-shadow:0 0 0 3px rgba(59,130,246,.2)!important;
}}
.stTextInput input:disabled,.stTextArea textarea:disabled{{opacity:0.55!important;cursor:not-allowed!important;}}

/* ── Selectbox ─────────────────────────────────────────────────── */
[data-baseweb="select"]>div,[data-baseweb="select"] [role="combobox"],
[data-baseweb="select"] span,[data-baseweb="select"] [data-value]{{
    background-color:{ibg}!important;color:{fg}!important;border-color:{brd}!important;
}}
[data-baseweb="popover"] [role="option"]{{background-color:{cbg}!important;color:{fg}!important;}}
[data-baseweb="popover"] [aria-selected="true"]{{background-color:{mbg}!important;}}

/* ── File uploader ─────────────────────────────────────────────── */
[data-testid="stFileUploaderDropzone"]{{
    background-color:{ibg}!important;
    border:1.5px dashed {brd}!important;border-radius:10px!important;
    position:relative;
}}
/* Browse button */
[data-testid="stFileUploaderDropzone"] button,
[data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"]{{
    background:{browse_bg}!important;color:{browse_fg}!important;
    border:1.5px solid {browse_brd}!important;border-radius:7px!important;
    font-weight:600!important;font-size:0.82rem!important;
}}
[data-testid="stFileUploaderDropzone"] button:hover,
[data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"]:hover{{
    background:{browse_hov}!important;border-color:{acc}!important;color:{acc}!important;
}}
/* Hide native English drag-drop text; show translated via ::after */
[data-testid="stFileUploaderDropzoneInstructions"] p,
[data-testid="stFileUploaderDropzone"] section>p,
[data-testid="stFileUploaderDropzone"]>section>div>p{{
    font-size:0!important;line-height:0!important;color:transparent!important;
    position:relative;display:block;
}}
[data-testid="stFileUploaderDropzoneInstructions"] p:first-of-type::after,
[data-testid="stFileUploaderDropzone"] section>p:first-of-type::after,
[data-testid="stFileUploaderDropzone"]>section>div>p:first-of-type::after{{
    content:'{drag_txt}';
    font-size:0.82rem!important;color:{sub}!important;display:block;
    line-height:1.4!important;
}}
/* Hide native file limit text */
[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderDropzoneInstructions"] small{{
    font-size:0!important;line-height:0!important;color:transparent!important;
    display:block;position:relative;
}}
[data-testid="stFileUploaderDropzone"] small::after,
[data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderDropzoneInstructions"] small::after{{
    content:'{limit_txt}';
    font-size:0.72rem!important;color:{sub}!important;
    display:block;line-height:1.4!important;
}}

/* ── Buttons — default accent ───────────────────────────────────── */
.stButton>button{{
    background:linear-gradient(135deg,{acc2},{acc})!important;
    color:white!important;border:none!important;border-radius:8px!important;
    font-weight:600!important;font-size:0.86rem!important;
    transition:all .18s;padding:0.42rem 0.9rem!important;
}}
.stButton>button:hover{{transform:translateY(-1px);box-shadow:0 4px 14px rgba(59,130,246,.4)!important;}}
.stButton>button:disabled{{
    background:{mbg}!important;color:{sub}!important;
    opacity:0.6!important;transform:none!important;box-shadow:none!important;
}}
/* Sidebar ghost buttons */
[data-testid="stSidebar"] .stButton>button{{
    background:{sb_nav_bg}!important;color:{sb_nav_fg}!important;
    border:1px solid {brd}!important;border-radius:7px!important;
    font-weight:500!important;font-size:0.78rem!important;
    box-shadow:none!important;transform:none!important;padding:0.3rem 0.6rem!important;
}}
[data-testid="stSidebar"] .stButton>button:hover{{
    background:{sb_nav_hov}!important;color:{fg}!important;
    border-color:{acc}!important;transform:none!important;box-shadow:none!important;
}}
/* Sidebar toggle */
[data-testid="stSidebar"] [data-testid="stToggle"] span,
[data-testid="stToggle"] label span{{color:{fg}!important;}}
/* Toggle track — visible in light mode */
[data-testid="stToggle"] [role="switch"]{{
    background-color:{"#334155" if dark else "#cbd5e1"}!important;
    border:2px solid {"#475569" if dark else "#94a3b8"}!important;
}}
[data-testid="stToggle"] [role="switch"][aria-checked="true"]{{
    background-color:{acc}!important;
    border-color:{acc}!important;
}}

/* Form submit */
[data-testid="stFormSubmitButton"]>button{{
    background:linear-gradient(135deg,{acc2},{acc})!important;
    color:white!important;border:none!important;border-radius:8px!important;
    font-weight:700!important;width:100%;font-size:0.92rem!important;
}}
/* Download */
.stDownloadButton>button{{
    background:{btn_dl_bg}!important;color:white!important;
    border:none!important;border-radius:8px!important;font-weight:600!important;
}}
.stDownloadButton>button:hover{{background:{btn_dl_hov}!important;}}

/* ── Tabs — equal width, fixed height ────────────────────────────── */
[data-baseweb="tab-list"]{{
    background:{"#1a2a42" if dark else "#e8f0fe"}!important;
    border-radius:12px!important;
    padding:5px!important;
    display:flex!important;
    gap:4px!important;
}}
[data-baseweb="tab"]{{
    flex:1!important;
    min-width:0!important;
    background:{"transparent" if dark else "#dbeafe"}!important;
    color:{"#94a3b8" if dark else "#1e40af"}!important;
    border-radius:9px!important;
    font-size:0.82rem!important;
    font-weight:{"500" if dark else "600"}!important;
    text-align:center!important;
    justify-content:center!important;
    white-space:nowrap!important;
    overflow:hidden!important;
    text-overflow:ellipsis!important;
    padding:7px 6px!important;
    border:{"none" if dark else "1px solid #bfdbfe"}!important;
    transition:background .15s,color .15s!important;
    cursor:pointer!important;
}}
[data-baseweb="tab"]:hover{{
    background:{"#243551" if dark else "#bfdbfe"}!important;
    color:{"#f1f5f9" if dark else "#1e3a8a"}!important;
}}
[aria-selected="true"][data-baseweb="tab"]{{
    background:{"#1e293b" if dark else "#2563eb"}!important;
    color:{"#60a5fa" if dark else "#ffffff"}!important;
    font-weight:700!important;
    border:{"1px solid #334155" if dark else "1px solid #2563eb"}!important;
    box-shadow:{"none" if dark else "0 2px 8px rgba(37,99,235,.25)"}!important;
}}

/* Expander */
[data-testid="stExpander"]{{background:{cbg}!important;border:1px solid {brd}!important;border-radius:10px!important;}}
[data-testid="stExpander"] summary *{{color:{acc}!important;font-weight:600!important;}}
[data-testid="stExpander"] summary svg{{fill:{acc}!important;}}
[data-testid="stExpander"] [data-testid="stExpanderDetails"] *{{color:{fg}!important;}}
[data-testid="stExpander"] [data-testid="stExpanderDetails"] .qa-q{{color:{fg}!important;}}
[data-testid="stExpander"] [data-testid="stExpanderDetails"] .qa-a{{color:{hint_fg}!important;}}
[data-testid="stExpander"] [data-testid="stExpanderDetails"] .hint-box *{{color:{hint_fg}!important;}}

/* DataFrame */
[data-testid="stDataFrame"] th{{background:{mbg}!important;color:{fg}!important;}}
[data-testid="stDataFrame"] td{{color:{fg}!important;background:{cbg}!important;}}

/* Progress */
[data-testid="stProgressBar"]>div{{background:{mbg}!important;}}
[data-testid="stProgressBar"]>div>div{{background:linear-gradient(90deg,{acc2},{acc})!important;}}
[data-testid="stProgressBar"] p{{color:{fg}!important;font-size:0.82rem!important;}}
hr{{border-color:{brd}!important;}}
code,pre{{background:{code_bg}!important;color:{'#93c5fd' if dark else '#1e40af'}!important;border-radius:6px!important;padding:2px 6px!important;}}
[data-testid="stAlert"]{{border-radius:10px!important;}}
[data-testid="stAlert"] p{{color:{fg}!important;}}
[data-testid="stMetricValue"],[data-testid="stMetricLabel"]{{color:{fg}!important;}}

/* ─── Custom HTML components ─────────────────────────────────────── */
.mc{{background:{mbg};border:1px solid {brd};border-radius:12px;padding:12px 14px;text-align:center;}}
.mv{{font-size:1.6rem;font-weight:800;color:{met}!important;}}
.ml{{font-size:0.72rem;color:{sub}!important;margin-top:2px;}}
.pv{{background:{preview_bg}!important;border:1px solid {brd};border-radius:12px;padding:20px 24px;font-family:'Times New Roman',serif;line-height:1.8;}}
.pv,.pv p,.pv div,.pv span,.pv b{{color:{fg}!important;}}
.pt{{font-size:1.15rem;font-weight:700;text-align:center;color:{fg}!important;margin-bottom:6px;}}
.pa{{text-align:center;color:{sub}!important;font-size:0.86rem;margin-bottom:8px;}}
.ph{{font-weight:700;color:{acc}!important;border-bottom:1px solid {brd};padding-bottom:3px;margin-top:14px;font-size:1rem;}}
.pab{{background:{mbg};border-left:4px solid {acc};padding:10px 14px;border-radius:0 8px 8px 0;margin:10px 0;font-size:0.86rem;color:{sub}!important;}}
.ri{{background:{cbg};border:1px solid {brd};border-radius:8px;padding:9px 13px;margin-bottom:5px;font-size:0.83rem;color:{fg}!important;}}
.rn{{color:{acc}!important;font-weight:700;}}
.lock-banner{{background:{lock_bg};border:1px solid {lock_brd};border-radius:8px;
    padding:8px 14px;margin:8px 0;font-size:0.84rem;color:{lock_fg}!important;
    display:flex;align-items:center;gap:8px;}}
.lock-banner *{{color:{lock_fg}!important;}}
.hint-box{{background:{hint_bg};border:1px solid {hint_brd};border-radius:10px;
    padding:12px 14px;margin:6px 0;font-size:0.83rem;color:{hint_fg}!important;line-height:1.65;}}
.hint-box *{{color:{hint_fg}!important;}}
.hint-lbl{{font-size:0.68rem;font-weight:700;color:{acc}!important;text-transform:uppercase;letter-spacing:.05em;margin-bottom:4px;}}
.qa-q{{font-weight:600;color:{fg}!important;font-size:0.85rem;margin-bottom:4px;}}
.qa-a{{font-size:0.83rem;color:{hint_fg}!important;line-height:1.6;padding-left:10px;border-left:2px solid {acc};}}
.ins-row{{background:{insert_bg};border-radius:8px;padding:6px 10px;margin:6px 0;display:flex;flex-wrap:wrap;gap:6px;align-items:center;}}
.ins-lbl{{font-size:0.7rem;color:{insert_fg}!important;font-weight:600;}}
.obj-ok{{background:{obj_ok};border-radius:7px;padding:5px 10px;font-size:0.8rem;color:{obj_ok_fg}!important;margin:2px 0;}}
.obj-miss{{background:{obj_miss};border-radius:7px;padding:5px 10px;font-size:0.8rem;color:{obj_miss_fg}!important;margin:2px 0;}}
.sec-hd{{font-size:1.05rem;font-weight:700;color:{fg}!important;margin-bottom:6px;}}
.qs{{font-size:0.7rem;padding:2px 7px;border-radius:6px;font-weight:700;display:inline-block;margin:1px;}}
.ai-badge{{background:{ai_bg};border:1px solid {ai_brd};border-radius:8px;
    padding:6px 10px;font-size:0.76rem;color:{ai_fg}!important;
    display:flex;align-items:center;gap:6px;margin-top:4px;}}
.ai-badge *{{color:{ai_fg}!important;}}
.ai-badge a{{color:{ai_fg}!important;text-decoration:underline;}}
.sb-sec-item{{font-size:0.76rem;padding:3px 4px;border-radius:5px;color:{sub}!important;line-height:1.4;}}
.sb-sec-done{{color:{fg}!important;font-weight:500;}}
footer{{visibility:hidden;}}#MainMenu{{visibility:hidden;}}
</style>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
#  GUEST BANNER + LOGIN PROMPT
# ═══════════════════════════════════════════════════════════════════════════
def guest_banner():
    if not is_logged_in():
        c1, c2 = st.columns([5, 1])
        c1.markdown(f'<div class="lock-banner">🔒 {t("guest_msg")}</div>',
                    unsafe_allow_html=True)
        if c2.button(t("login_link"), key="gb_login"):
            st.session_state.page = "auth"; st.rerun()

def login_prompt_button(key_suffix: str):
    """Show an actionable 'Sign in to edit' button for guests."""
    if not is_logged_in():
        if st.button(f"🔓 {t('login_link')}", key=f"lp_{key_suffix}"):
            st.warning(t("login_required"))
            st.session_state.page = "auth"; st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
#  INLINE NAV  (equal-size buttons with stats in brackets)
# ═══════════════════════════════════════════════════════════════════════════
ALL_PAGES = [
    ("info",    "📄", "nav_info"),
    ("sections","✍️","nav_sec"),
    ("media",   "🖼️","nav_media"),
    ("refs",    "📑","nav_refs"),
    ("generate","🚀","nav_gen"),
    ("thesis",  "🎓","nav_thesis"),
    ("agent",   "🤖","nav_agent"),
    ("settings","⚙️","nav_set"),
]

def _nav_stat(key: str) -> str:
    """Return compact stat string shown in nav button brackets."""
    if key == "info":
        w = wc(st.session_state.get("abstract","") + " " + st.session_state.get("art_title",""))
        return f" ({w})" if w else ""
    if key == "sections":
        done = sum(1 for k in active_sections() if st.session_state.get(k,"").strip())
        tot  = len(active_sections())
        return f" ({done}/{tot})"
    if key == "media":
        n = len(st.session_state.figures) + len(st.session_state.tables) + len(st.session_state.formulas)
        return f" ({n})" if n else ""
    if key == "refs":
        n = len(st.session_state.refs)
        return f" ({n})" if n else ""
    if key == "generate":
        score, _ = article_readiness()
        return f" ({score}%)"
    if key == "thesis":
        done = sum(1 for k in THESIS_SECTIONS if st.session_state.get(k,"").strip())
        return f" ({done})" if done else ""
    return ""

def inline_nav():
    dark   = st.session_state.dark
    acc_c  = "#3b82f6" if dark else "#2563eb"
    acc2_c = "#2563eb" if dark else "#1e40af"
    # Light mode matches the tab pill style from screenshot
    ibtn_bg  = "#1e293b"  if dark else "#dbeafe"
    ibtn_fg  = "#94a3b8"  if dark else "#1e40af"
    ibtn_brd = "#334155"  if dark else "#bfdbfe"
    ihov_bg  = "#243551"  if dark else "#bfdbfe"
    ihov_fg  = "#f1f5f9"  if dark else "#1e3a8a"
    top_brd  = "#334155"  if dark else "#e8f0fe"
    act_bg   = f"linear-gradient(150deg,{acc2_c},{acc_c})" if dark else "#2563eb"
    act_shad = "rgba(59,130,246,.3)" if dark else "rgba(37,99,235,.25)"
    n = len(ALL_PAGES)

    st.markdown(f"""<style>
/* ── Nav row ───────────────────────────────────────────────────── */
[data-testid="stHorizontalBlock"]:has(>[data-testid="column"]:nth-child({n})) {{
    border-top:1px solid {top_brd};
    border-bottom:1px solid {top_brd};
    padding:6px 0;
    margin-bottom:14px;
    gap:4px;
    background:{"transparent" if dark else "#e8f0fe"};
    border-radius:{"0" if dark else "14px"};
}}
/* ── Shared shape ─────────────────────────────────────────────── */
[data-testid="stHorizontalBlock"]:has(>[data-testid="column"]:nth-child({n}))
.stButton>button {{
    width:100%!important;
    height:56px!important;
    min-height:56px!important;
    max-height:56px!important;
    display:flex!important;
    flex-direction:column!important;
    align-items:center!important;
    justify-content:center!important;
    text-align:center!important;
    white-space:normal!important;
    word-break:break-word!important;
    overflow:hidden!important;
    font-size:0.72rem!important;
    line-height:1.18!important;
    padding:4px 3px!important;
    border-radius:10px!important;
    box-shadow:none!important;
    transform:none!important;
    font-weight:{"500" if dark else "600"}!important;
    transition:background .14s,color .14s,border-color .14s!important;
}}
/* ── Inactive ─────────────────────────────────────────────────── */
[data-testid="stHorizontalBlock"]:has(>[data-testid="column"]:nth-child({n}))
[data-testid="baseButton-secondary"] {{
    background:{ibtn_bg}!important;
    color:{ibtn_fg}!important;
    border:1px solid {ibtn_brd}!important;
}}
[data-testid="stHorizontalBlock"]:has(>[data-testid="column"]:nth-child({n}))
[data-testid="baseButton-secondary"]:hover {{
    background:{ihov_bg}!important;
    color:{ihov_fg}!important;
    border-color:{acc_c}!important;
    transform:none!important;
    box-shadow:none!important;
}}
/* ── Active ───────────────────────────────────────────────────── */
[data-testid="stHorizontalBlock"]:has(>[data-testid="column"]:nth-child({n}))
[data-testid="baseButton-primary"] {{
    background:{act_bg}!important;
    color:#ffffff!important;
    border:{"none" if dark else "1px solid "+acc_c}!important;
    font-weight:700!important;
    box-shadow:0 2px 10px {act_shad}!important;
}}
[data-testid="stHorizontalBlock"]:has(>[data-testid="column"]:nth-child({n}))
[data-testid="baseButton-primary"]:hover {{
    transform:none!important;
    box-shadow:0 3px 12px {act_shad}!important;
}}
</style>""", unsafe_allow_html=True)

    cols = st.columns(n)
    for col, (key, icon, tk) in zip(cols, ALL_PAGES):
        stat      = _nav_stat(key)
        label     = f"{icon} {t(tk)}{stat}"
        is_active = st.session_state.page == key
        btn_type  = "primary" if is_active else "secondary"
        if col.button(label, key=f"nav_{key}",
                      use_container_width=True, type=btn_type):
            st.session_state.page = key; st.rerun()

# stats_row removed — stats now shown in inline_nav brackets

# ═══════════════════════════════════════════════════════════════════════════
#  ARTICLE READINESS
# ═══════════════════════════════════════════════════════════════════════════
def article_readiness() -> tuple:
    score   = 0
    missing = []
    for val, key, pts in [
        (st.session_state.art_title, "title", 15),
        (st.session_state.authors,   "authors", 10),
        (st.session_state.abstract,  "abstract", 15),
        (st.session_state.keywords,  "keywords", 5),
    ]:
        if val.strip(): score += pts
        else: missing.append(key)
    req = [s for s in active_sections() if s in
           ["intro","materials_methods","results","conclusion"]]
    for s in req:
        if st.session_state.get(s,"").strip(): score += 10
        else: missing.append(s)
    if st.session_state.refs: score += 5
    else: missing.append("references")
    return min(score, 100), missing

# ═══════════════════════════════════════════════════════════════════════════
#  OBJECTIVES CHECKER  (interactive — live update from intro text)
# ═══════════════════════════════════════════════════════════════════════════
def objectives_checker(intro_override: str = ""):
    lk  = _lang_key()
    txt = (intro_override or st.session_state.get("intro","")).lower()

    titles = {
        "kz": "🎯 Мақсат пен міндеттерді тексеру",
        "ru": "🎯 Проверка цели и задач",
        "en": "🎯 Objectives & Tasks Check",
    }
    st.markdown(f'<div class="hint-lbl">{titles[lk]}</div>', unsafe_allow_html=True)

    if not txt.strip():
        st.markdown(
            f'<div class="hint-box" style="opacity:.6;">'
            f'{"✍️ Кіріспені жазыңыз — тексеріс өзіндік іске қосылады." if lk=="kz" else "✍️ Напишите введение — проверка запустится автоматически." if lk=="ru" else "✍️ Write the Introduction — the check will run automatically."}'
            f'</div>', unsafe_allow_html=True)
        return

    aim_kw  = t("aim_kw")
    task_kw = t("task_kw")
    aim_found  = any(kw in txt for kw in aim_kw)
    task_found = any(kw in txt for kw in task_kw)
    word_cnt   = wc(txt)
    has_len    = word_cnt >= 150
    has_hypothesis = any(w in txt for w in
        ["гипотеза","болжам","hypothesis","предположение","предполагается"])
    has_novelty    = any(w in txt for w in
        ["жаңалығы","новизна","novelty","new approach","впервые"])

    checks = [
        (aim_found,  {"kz":"Зерттеу мақсаты анықталған","ru":"Цель исследования указана","en":"Research objective stated"}[lk]),
        (task_found, {"kz":"Міндеттер тізімі бар","ru":"Задачи перечислены","en":"Research tasks listed"}[lk]),
        (has_len,    {"kz":f"Жеткілікті көлем ({word_cnt} сөз ≥ 150)","ru":f"Достаточный объём ({word_cnt} слов ≥ 150)","en":f"Sufficient length ({word_cnt} words ≥ 150)"}[lk]),
        (has_hypothesis,{"kz":"Гипотеза / болжам бар","ru":"Гипотеза присутствует","en":"Hypothesis mentioned"}[lk]),
        (has_novelty,   {"kz":"Зерттеудің жаңалығы көрсетілген","ru":"Научная новизна отмечена","en":"Scientific novelty mentioned"}[lk]),
    ]

    score = sum(1 for ok, _ in checks if ok)
    total = len(checks)
    pct   = int(score / total * 100)
    bar_color = "#22c55e" if pct >= 80 else ("#f59e0b" if pct >= 40 else "#ef4444")
    c_    = _colors()

    st.markdown(
        f'<div style="background:{c_["mbg"]};'
        f'border-radius:8px;padding:8px 12px;margin:4px 0;">'
        f'<div style="display:flex;align-items:center;gap:8px;margin-bottom:6px;">'
        f'<div style="flex:1;height:6px;background:{c_["brd"]};border-radius:3px;">'
        f'<div style="width:{pct}%;height:100%;background:{bar_color};border-radius:3px;transition:width .4s;"></div>'
        f'</div><span style="font-size:0.76rem;font-weight:700;color:{bar_color};">{pct}%</span></div>',
        unsafe_allow_html=True)

    for ok, label in checks:
        icon  = "✅" if ok else "⚪"
        cls   = "obj-ok" if ok else "obj-miss"
        st.markdown(f'<div class="{cls}" style="margin:2px 0;">{icon} {label}</div>',
                    unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Actionable tips for missing items
    missing_tips = []
    if not aim_found:
        missing_tips.append({"kz":"'Зерттеудің мақсаты — ...' деген тіркесімен мақсатты нақты жазыңыз.",
                              "ru":"Добавьте фразу: 'Целью работы является...'",
                              "en":"Add the phrase: 'The objective of this study is to...'"}[lk])
    if not task_found:
        missing_tips.append({"kz":"Нөмірленген тізім түрінде 3–5 міндетті көрсетіңіз.",
                              "ru":"Перечислите 3–5 задач нумерованным списком.",
                              "en":"List 3–5 tasks as a numbered list."}[lk])
    if not has_len:
        need = 150 - word_cnt
        missing_tips.append({"kz":f"Тағы {need} сөз қосыңыз (мин. 150 сөз).",
                              "ru":f"Добавьте ещё {need} слов (мин. 150).",
                              "en":f"Add {need} more words (min. 150)."}[lk])

    if missing_tips:
        tips_lbl = {"kz":"💡 Ұсыныстар:","ru":"💡 Рекомендации:","en":"💡 Suggestions:"}[lk]
        st.markdown(f'<div class="hint-box" style="margin-top:6px;"><div class="hint-lbl">{tips_lbl}</div>',
                    unsafe_allow_html=True)
        for tip in missing_tips:
            st.markdown(f'<div style="margin:3px 0;">→ {tip}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
#  AI INTERACTIVE HINT PANEL  (with live Anthropic API call)
# ═══════════════════════════════════════════════════════════════════════════
def _call_anthropic(prompt: str, system: str = "") -> str:
    """Call Anthropic API; return response text or error string."""
    api_key = st.session_state.get("anthropic_api_key","")
    if not api_key:
        return ""
    if not REQ_OK:
        return "⚠️ pip install requests"
    try:
        body = {
            "model": "claude-haiku-4-5-20251001",
            "max_tokens": 400,
            "messages": [{"role":"user","content": prompt}],
        }
        if system:
            body["system"] = system
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json=body,
            timeout=20,
        )
        if resp.status_code == 200:
            return resp.json()["content"][0]["text"]
        return f"⚠️ API {resp.status_code}: {resp.json().get('error',{}).get('message','')}"
    except Exception as e:
        return f"⚠️ {e}"

def ai_hint_panel(section_key: str):
    data = SEC_HINTS.get(section_key)
    if not data: return
    lk          = _lang_key()
    tip         = data["tip"].get(lk, "")
    has_ai_key  = bool(st.session_state.get("anthropic_api_key",""))
    cur_text    = st.session_state.get(section_key,"")

    # AI active badge
    if has_ai_key:
        ai_lbl = {"kz":"✦ ЖИ белсенді","ru":"✦ ИИ активен","en":"✦ AI active"}[lk]
        st.markdown(f'<div class="ai-badge">🤖 {ai_lbl} — Claude Haiku 4.5</div>',
                    unsafe_allow_html=True)

    # Static tip
    if tip:
        st.markdown(f'<div class="hint-box"><div class="hint-lbl">✦ Smart Article</div>{tip}</div>',
                    unsafe_allow_html=True)

    # Interactive Q&A
    qa = data.get("qa",[])
    if qa:
        qa_lbl = {"kz":"💬 Сұрақ-жауап","ru":"💬 Вопросы и ответы","en":"💬 Q&A"}[lk]
        with st.expander(qa_lbl):
            for item in qa:
                q = item["q"].get(lk,"")
                a = item["a"].get(lk,"")
                st.markdown(f'<div class="qa-q">❓ {q}</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="qa-a">{a}</div>', unsafe_allow_html=True)
                st.write("")

    # ── AI interactive button ─────────────────────────────────────
    if has_ai_key and cur_text.strip():
        ai_btn_lbl = {
            "kz": "✦ ЖИ кеңесін алу",
            "ru": "✦ Получить совет ИИ",
            "en": "✦ Get AI feedback",
        }[lk]
        btn_key = f"ai_ask_{section_key}"
        result_key = f"ai_result_{section_key}"

        if st.button(ai_btn_lbl, key=btn_key, use_container_width=False):
            sec_display = sec_name(section_key)
            sys_prompt  = {
                "kz": (f"Сіз ғылыми жазуға арналған көмекші боласыз. "
                       f"Мақаланың '{sec_display}' бөлімін қысқаша талдап, 3 нақты жақсарту ұсынысын беріңіз. "
                       f"Мәтінді өзіңіз жазбаңыз — тек нұсқаулар беріңіз. Қазақша жазыңыз."),
                "ru": (f"Вы помощник по научному письму. "
                       f"Кратко проанализируйте раздел '{sec_display}' и дайте 3 конкретных совета по улучшению. "
                       f"Не переписывайте текст — только рекомендации. Отвечайте на русском."),
                "en": (f"You are a scientific writing assistant. "
                       f"Briefly analyse the '{sec_display}' section and give 3 specific improvement tips. "
                       f"Do not rewrite the text — only give actionable suggestions. Reply in English."),
            }[lk]
            user_msg = f"{cur_text[:1500]}"
            with st.spinner("✦ AI..."):
                result = _call_anthropic(user_msg, system=sys_prompt)
            st.session_state[result_key] = result

        # Show cached result
        cached = st.session_state.get(result_key, "")
        if cached:
            ai_result_lbl = {"kz":"ЖИ кеңесі","ru":"Совет ИИ","en":"AI feedback"}[lk]
            st.markdown(
                f'<div class="hint-box" style="margin-top:6px;">'
                f'<div class="hint-lbl">✦ {ai_result_lbl}</div>'
                f'{cached}</div>',
                unsafe_allow_html=True)
            if st.button({"kz":"✕ Жабу","ru":"✕ Закрыть","en":"✕ Clear"}[lk],
                         key=f"ai_clear_{section_key}"):
                st.session_state[result_key] = ""
                st.rerun()

    elif not has_ai_key:
        no_key_lbl = {
            "kz": "🤖 ЖИ кеңесін алу үшін API кілтін енгізіңіз",
            "ru": "🤖 Введите API ключ для получения советов ИИ",
            "en": "🤖 Enter API key to get AI feedback",
        }[lk]
        c_    = _colors()
        lnk_c = "#93c5fd" if st.session_state.dark else "#1d4ed8"
        st.markdown(
            f'<div style="font-size:0.76rem;color:{c_["muted"]};margin:4px 0;">{no_key_lbl} '
            f'<a href="https://console.anthropic.com/settings/keys" target="_blank" '
            f'style="color:{lnk_c};">→ console.anthropic.com</a>'
            f'</div>',
            unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
#  DYNAMIC INSERT BUTTONS
# ═══════════════════════════════════════════════════════════════════════════
def insert_buttons(section_key: str):
    """Render compact insert buttons for figures, tables, formulas."""
    if not is_logged_in(): return
    figs  = st.session_state.figures
    tbls  = st.session_state.tables
    forms = st.session_state.formulas
    if not (figs or tbls or forms): return

    st.markdown(f'<div class="ins-row"><span class="ins-lbl">↩ Insert:</span></div>',
                unsafe_allow_html=True)
    cols = st.columns(min(len(figs)+len(tbls)+len(forms)+1, 6))
    cidx = 0

    for fig in figs:
        label = f"{t('fig_lbl')} {fig['number']}"
        if cidx < len(cols):
            if cols[cidx].button(f"🖼️ {label}", key=f"ins_fig_{section_key}_{fig['number']}",
                                  use_container_width=True):
                cur = st.session_state.get(section_key,"")
                st.session_state[section_key] = cur + f" ({t('fig_lbl')} {fig['number']})"
                st.rerun()
        cidx += 1

    for tbl in tbls:
        label = f"{t('tbl_lbl')} {tbl['number']}"
        if cidx < len(cols):
            if cols[cidx].button(f"📊 {label}", key=f"ins_tbl_{section_key}_{tbl['number']}",
                                  use_container_width=True):
                cur = st.session_state.get(section_key,"")
                st.session_state[section_key] = cur + f" ({t('tbl_lbl')} {tbl['number']})"
                st.rerun()
        cidx += 1

    for frm in forms:
        label = f"{t('form_lbl')} {frm['number']}"
        if cidx < len(cols):
            if cols[cidx].button(f"🧮 {label}", key=f"ins_frm_{section_key}_{frm['number']}",
                                  use_container_width=True):
                cur = st.session_state.get(section_key,"")
                st.session_state[section_key] = cur + f" ({t('form_lbl')} ({frm['number']}))"
                st.rerun()
        cidx += 1

# ═══════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════
def sidebar():
    with st.sidebar:
        dark = st.session_state.dark
        lk   = _lang_key()
        fg_  = "#f1f5f9" if dark else "#0d1f3c"
        sub_ = "#94a3b8" if dark else "#475569"
        acc_ = "#3b82f6" if dark else "#1d4ed8"
        muted_ = "#64748b" if dark else "#6b7280"

        # Logo
        st.markdown(
            f'<div style="font-size:1.25rem;font-weight:900;color:{acc_};line-height:1.2;">📝 Smart Article</div>'
            f'<div style="font-size:0.65rem;color:{muted_};margin-bottom:8px;">Research Writing Platform</div>',
            unsafe_allow_html=True)

        # ── Theme toggle ─────────────────────────────────────────
        c1, c2 = st.columns(2)
        if c1.button(t("dark_mode"),  key="sb_dark",  use_container_width=True):
            st.session_state.dark = True;  st.rerun()
        if c2.button(t("light_mode"), key="sb_light", use_container_width=True):
            st.session_state.dark = False; st.rerun()

        st.divider()

        # ── Language ─────────────────────────────────────────────
        langs = ["🇰🇿 Қазақша","🇷🇺 Русский","🇬🇧 English"]
        sel_lang = st.selectbox("🌍", langs,
                                index=langs.index(st.session_state.lang),
                                label_visibility="collapsed", key="lang_sb")
        if sel_lang != st.session_state.lang:
            st.session_state.lang = sel_lang; st.rerun()

        # ── Template ─────────────────────────────────────────────
        tpl_keys = list(TEMPLATES.keys())
        sel_tpl  = st.selectbox("📋", tpl_keys,
                                index=tpl_keys.index(st.session_state.template)
                                if st.session_state.template in tpl_keys else 0,
                                label_visibility="collapsed", key="tpl_sb")
        if sel_tpl != st.session_state.template:
            st.session_state.template = sel_tpl; st.rerun()

        st.divider()

        # ── Article readiness ─────────────────────────────────────
        score, _ = article_readiness()
        st.progress(score / 100, text=f"{t('readiness')}: {score}%")

        st.divider()

        # ── AI activation ────────────────────────────────────────
        ai_key = st.session_state.get("anthropic_api_key","")
        lk_ai  = {"kz":"✦ ЖИ белсендіру","ru":"✦ Активация ИИ","en":"✦ Enable AI"}[lk]
        if not ai_key:
            st.markdown(
                f'<div class="ai-badge">🤖 <span>'
                f'<a href="https://console.anthropic.com/settings/keys" target="_blank">'
                f'{lk_ai}</a></span></div>',
                unsafe_allow_html=True)
            with st.expander({"kz":"API кілтін енгізу","ru":"Ввести API ключ","en":"Enter API key"}[lk]):
                new_key = st.text_input("Anthropic API Key",
                                        placeholder="sk-ant-...",
                                        type="password", key="ai_key_input")
                if st.button({"kz":"Сақтау","ru":"Сохранить","en":"Save"}[lk],
                             key="save_ai_key"):
                    st.session_state.anthropic_api_key = new_key
                    st.success("✅")
                    st.rerun()
        else:
            masked = f"sk-ant-...{ai_key[-6:]}" if len(ai_key) > 10 else "***"
            st.markdown(
                f'<div class="ai-badge">🤖 AI: {masked}</div>',
                unsafe_allow_html=True)
            if st.button({"kz":"Кілтті өзгерту","ru":"Сменить ключ","en":"Change key"}[lk],
                         key="clear_ai_key"):
                st.session_state.anthropic_api_key = ""
                st.rerun()

        st.divider()

        # ── User ─────────────────────────────────────────────────
        if is_logged_in():
            ud = st.session_state.user_data
            name_str = ud.get("name", st.session_state.username)
            role_str = ud.get("role","")
            st.markdown(
                f'<div style="font-size:0.76rem;color:{sub_};">👤 '
                f'<b style="color:{fg_};">{name_str}</b><br>'
                f'<span style="font-size:0.66rem;color:{muted_};">{role_str}</span></div>',
                unsafe_allow_html=True)
            if st.button(f"🚪 {t('logout')}", use_container_width=True, key="sb_logout"):
                add_log("logout", st.session_state.username)
                st.session_state.logged_in = False; st.rerun()
        else:
            if st.button(f"🔑 {t('sign_in')} / {t('register')}", use_container_width=True,
                         key="sb_signin"):
                st.session_state.page = "auth"; st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: AUTH
# ═══════════════════════════════════════════════════════════════════════════
def pg_auth():
    if is_logged_in():
        st.session_state.page = "info"; st.rerun(); return
    st.markdown(f"## 🔑 {t('sign_in')} / {t('register')}")
    inline_nav()
    _, mid, _ = st.columns([1,1.3,1])
    with mid:
        tab_l, tab_r = st.tabs([t("sign_in"), t("register")])
        with tab_l:
            with st.form("lf"):
                u = st.text_input(f"👤 {t('username')}", placeholder="admin")
                p = st.text_input(f"🔑 {t('password')}", type="password")
                ok = st.form_submit_button(t("login_btn"), use_container_width=True)
            if ok:
                s, ud = do_login(u.strip(), p)
                if s:
                    st.session_state.logged_in = True
                    st.session_state.username  = u.strip()
                    st.session_state.user_data = ud
                    st.session_state.page      = "info"; st.rerun()
                else:
                    st.error(f"❌ {t('login_err')}")
            st.caption("Demo: **admin** / **admin123**")
        with tab_r:
            with st.form("rf"):
                ru = st.text_input(f"👤 {t('username')}", key="ru")
                rn = st.text_input(f"🙍 {t('full_name')}", key="rn")
                re_ = st.text_input(f"📧 {t('email')}", key="re")
                rr  = st.selectbox(f"🎓 {t('role')}", t("roles"))
                rp  = st.text_input(f"🔑 {t('password')}", type="password", key="rp")
                rp2 = st.text_input(f"🔑 {t('confirm_pw')}", type="password", key="rp2")
                rok = st.form_submit_button(t("reg_btn"), use_container_width=True)
            if rok:
                if rp != rp2: st.error(f"❌ {t('pw_mismatch')}")
                else:
                    ok2, code = do_register(ru.strip(), re_.strip(), rp, rn.strip(), rr)
                    if ok2: st.success(t("reg_ok"))
                    else:
                        err = {"uname_short":t("uname_short"),"pw_short":t("pw_short"),
                               "username_taken":t("username_taken"),"email_taken":t("email_taken")}
                        st.error(f"❌ {err.get(code,code)}")

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: INFO
# ═══════════════════════════════════════════════════════════════════════════
def pg_info():
    st.markdown(f"## 📄 {t('nav_info')}")
    guest_banner(); inline_nav(); st.write("")
    dis  = not is_logged_in()
    tpl  = current_template()
    L, R = st.columns(2)

    with L:
        if tpl.get("irsti_label"):
            ca, cb = st.columns(2)
            v = ca.text_input(f"🔢 {tpl['irsti_label']}", value=st.session_state.irsti,
                              placeholder=tpl.get("irsti_hint",""), disabled=dis)
            if not dis: st.session_state.irsti = v
            v2 = cb.text_input("📚 Section", value=st.session_state.section_field, disabled=dis)
            if not dis: st.session_state.section_field = v2
            if dis:
                login_prompt_button("irsti")

        atp = t("art_types")
        vi  = atp.index(st.session_state.art_type) if st.session_state.art_type in atp else 0
        v   = st.selectbox(f"📌 {t('art_type_f')}", atp, index=vi, disabled=dis)
        if not dis: st.session_state.art_type = v

        v = st.text_input(f"📰 {t('title_f')}", value=st.session_state.art_title,
                          placeholder="Enter full article title", disabled=dis)
        if not dis: st.session_state.art_title = v
        if dis: login_prompt_button("title")

        v = st.text_area(f"👥 {t('authors_f')}", value=st.session_state.authors,
                         height=65, disabled=dis,
                         placeholder="Lastname F.1, Lastname F.2,* — use superscript numbers")
        if not dis: st.session_state.authors = v

        v = st.text_area(f"🏛️ {t('affil_f')}", value=st.session_state.affiliation,
                         height=65, disabled=dis,
                         placeholder="1 University, City, Country; email@domain.com")
        if not dis: st.session_state.affiliation = v

        v = st.text_input(f"📖 {t('journal_f')}", value=st.session_state.journal,
                          disabled=dis, placeholder="e.g. Vestnik ENU")
        if not dis: st.session_state.journal = v

    with R:
        max_w = tpl.get("abstract_max",300)
        kw_ph = f"{tpl.get('keywords_count','3–10')} keywords; separated by semicolons"
        v = st.text_input(f"🏷️ {t('keywords_f')}", value=st.session_state.keywords,
                          placeholder=kw_ph, disabled=dis)
        if not dis: st.session_state.keywords = v

        v = st.text_area(f"📝 {t('abstract_f')}", value=st.session_state.abstract,
                         height=170, disabled=dis,
                         placeholder=f"Max {max_w} words.")
        if not dis: st.session_state.abstract = v
        w = wc(st.session_state.abstract)
        color = "#10b981" if w <= max_w else "#ef4444"
        st.markdown(f'<span style="color:{color};font-size:0.78rem;">📝 {t("word_count")}: {w}/{max_w}</span>',
                    unsafe_allow_html=True)

        if "ENU Journal (EN)" in st.session_state.template:
            v = st.text_area("📝 Abstract (KZ)", value=st.session_state.abstract_kz,
                             height=70, disabled=dis, key="abs_kz")
            if not dis: st.session_state.abstract_kz = v
            v = st.text_area("📝 Abstract (RU)", value=st.session_state.abstract_ru,
                             height=70, disabled=dis, key="abs_ru")
            if not dis: st.session_state.abstract_ru = v

        if dis: login_prompt_button("abstract")
        st.markdown("---")
        ab  = st.session_state.abstract
        c_  = _colors()
        st.markdown(
            f'<div class="pv" style="padding:14px;">'
            f'<div class="pt">{st.session_state.art_title or "—"}</div>'
            f'<div class="pa">{st.session_state.authors or "—"}</div>'
            f'<div class="pa" style="font-size:0.78rem;">{st.session_state.affiliation}</div>'
            f'{"<div class=pab><b>Abstract:</b> "+ab[:250]+( chr(8230) if len(ab)>250 else "")+"</div>" if ab else ""}'
            f'{"<div style=font-size:0.75rem;color:"+c_["muted"]+";margin-top:4px;><b>Keywords:</b> "+st.session_state.keywords+"</div>" if st.session_state.keywords else ""}'
            f'</div>', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: SECTIONS
# ═══════════════════════════════════════════════════════════════════════════
def pg_sections():
    st.markdown(f"## ✍️ {t('nav_sec')}")
    guest_banner(); inline_nav(); st.write("")
    dis  = not is_logged_in()
    secs = active_sections()

    if "_active_sec" not in st.session_state or st.session_state._active_sec not in secs:
        st.session_state._active_sec = secs[0]

    tab_labels = [f"{ALL_SECTIONS.get(k,{}).get('icon','•')} {sec_name(k)}" for k in secs]
    tabs       = st.tabs(tab_labels)

    for tab, key in zip(tabs, secs):
        with tab:
            st.session_state._active_sec = key
            st.markdown(f'<div class="sec-hd">{ALL_SECTIONS.get(key,{}).get("icon","•")} {sec_name(key)}</div>',
                        unsafe_allow_html=True)
            L, R = st.columns([3,2], gap="large")

            with L:
                # Upload helper
                up = st.file_uploader(f"TXT/DOCX", type=["docx","txt"],
                                      key=f"up_{key}", disabled=dis,
                                      label_visibility="collapsed")
                if up and not dis:
                    if up.name.endswith(".txt"):
                        st.session_state[key] = up.read().decode("utf-8", errors="ignore")
                    elif up.name.endswith(".docx") and DOCX_OK:
                        try:
                            doc = Document(BytesIO(up.read()))
                            st.session_state[key] = "\n".join(p.text for p in doc.paragraphs)
                        except: pass

                # Insert buttons
                insert_buttons(key)

                val = st.text_area(
                    sec_name(key), value=st.session_state.get(key,""),
                    height=340, key=f"ta_{key}", disabled=dis,
                    label_visibility="collapsed",
                    placeholder=f"{sec_name(key)}…")
                if not dis: st.session_state[key] = val
                st.caption(f"📝 {t('word_count')}: {wc(st.session_state.get(key,''))}")
                if dis: login_prompt_button(f"sec_{key}")

                # Objectives checker in intro — live, reads current textarea value
                if key == "intro" and is_logged_in():
                    st.markdown("---")
                    objectives_checker(intro_override=st.session_state.get("intro",""))

            with R:
                ai_hint_panel(key)
                st.markdown(f"**{t('preview_lbl')}**")
                c  = st.session_state.get(key,"")
                c_ = _colors()
                st.markdown(
                    f'<div class="pv" style="padding:12px;font-size:0.86rem;">'
                    f'<div class="ph">{sec_name(key)}</div>'
                    f'{"<p style=margin-top:8px;>"+c[:500]+( chr(8230) if len(c)>500 else "")+"</p>" if c else "<p style=color:"+c_["muted"]+";font-style:italic;>Empty…</p>"}'
                    f'</div>', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: MEDIA
# ═══════════════════════════════════════════════════════════════════════════
def pg_media():
    st.markdown(f"## 🖼️ {t('nav_media')}")
    guest_banner(); inline_nav(); st.write("")
    dis = not is_logged_in()
    tab_fig, tab_tbl, tab_frm = st.tabs(
        [f"🖼️ {t('w_figs')}", f"📊 {t('w_tbls')}", f"🧮 {t('w_forms')}"])

    # ── FIGURES ───────────────────────────────────────────────────
    with tab_fig:
        L, R = st.columns(2, gap="large")
        with L:
            st.markdown(f'<div class="sec-hd">➕ {t("add_fig")}</div>', unsafe_allow_html=True)
            ai_hint_panel("")  # generic hint
            with st.form("ff", clear_on_submit=True):
                fn  = st.text_input(f"№", placeholder="1", disabled=dis)
                fc  = st.text_input(t("caption_lbl"),
                                    placeholder=f"Figure 1. Map of study area", disabled=dis)
                fup = st.file_uploader(t("upload_img"),
                                       type=["png","jpg","jpeg","tif","svg"], disabled=dis)
                if st.form_submit_button(f"➕ {t('add_btn')}", use_container_width=True,
                                         disabled=dis) and fc:
                    st.session_state.figures.append({
                        "number": fn or str(len(st.session_state.figures)+1),
                        "caption": fc,
                        "image": fup.read() if fup else None,
                        "name": fup.name if fup else None,
                    })
                    st.success("✅"); st.rerun()
            if dis: login_prompt_button("fig")

        with R:
            st.markdown(f'<div class="sec-hd">📋 {t("w_figs")} ({len(st.session_state.figures)})</div>',
                        unsafe_allow_html=True)
            if not st.session_state.figures:
                st.info(t("no_figs"))
            else:
                for i, fig in enumerate(st.session_state.figures):
                    with st.expander(f"🖼️ {t('fig_lbl')} {fig['number']} — {fig['caption'][:50]}"):
                        if fig.get("image"): st.image(fig["image"], use_container_width=True)
                        st.caption(fig.get("name","—"))
                        if not dis and st.button(f"🗑️ {t('del_btn')}", key=f"df{i}"):
                            st.session_state.figures.pop(i); st.rerun()

    # ── TABLES ────────────────────────────────────────────────────
    with tab_tbl:
        L, R = st.columns(2, gap="large")
        with L:
            st.markdown(f'<div class="sec-hd">➕ {t("add_tbl")}</div>', unsafe_allow_html=True)
            with st.form("tf", clear_on_submit=True):
                tn = st.text_input("№", placeholder="1", disabled=dis)
                tc = st.text_input(t("caption_lbl"),
                                   placeholder="Table 1. Summary statistics", disabled=dis)
                td = st.text_area(t("csv_data"), height=110,
                                  placeholder="Col1,Col2,Col3\nVal1,Val2,Val3", disabled=dis)
                if st.form_submit_button(f"➕ {t('add_btn')}", use_container_width=True,
                                         disabled=dis) and tc:
                    st.session_state.tables.append({
                        "number": tn or str(len(st.session_state.tables)+1),
                        "caption": tc, "data": td,
                    })
                    st.success("✅"); st.rerun()
            if dis: login_prompt_button("tbl")

        with R:
            st.markdown(f'<div class="sec-hd">📋 {t("w_tbls")} ({len(st.session_state.tables)})</div>',
                        unsafe_allow_html=True)
            if not st.session_state.tables:
                st.info(t("no_tbls"))
            else:
                for i, tbl in enumerate(st.session_state.tables):
                    with st.expander(f"📊 {t('tbl_lbl')} {tbl['number']} — {tbl['caption'][:50]}"):
                        if tbl.get("data") and PD_OK:
                            try:
                                df = pd.read_csv(io.StringIO(tbl["data"]))
                                st.dataframe(df, use_container_width=True)
                            except: st.text(tbl["data"])
                        if not dis and st.button(f"🗑️ {t('del_btn')}", key=f"dt{i}"):
                            st.session_state.tables.pop(i); st.rerun()

    # ── FORMULAS ──────────────────────────────────────────────────
    with tab_frm:
        L, R = st.columns(2, gap="large")
        with L:
            st.markdown(f'<div class="sec-hd">➕ {t("add_form")}</div>', unsafe_allow_html=True)
            with st.form("formf", clear_on_submit=True):
                fnum   = st.text_input("№", placeholder="1", disabled=dis)
                flatex = st.text_input(t("latex_lbl"),
                                       placeholder=r"Q = \frac{1}{n} A R^{2/3} S^{1/2}",
                                       disabled=dis)
                fdesc  = st.text_area(t("desc_lbl"), height=60, disabled=dis,
                                      placeholder="Manning's equation")
                if st.form_submit_button(f"➕ {t('add_btn')}", use_container_width=True,
                                         disabled=dis) and flatex:
                    st.session_state.formulas.append({
                        "number": fnum or str(len(st.session_state.formulas)+1),
                        "latex": flatex, "desc": fdesc,
                    })
                    st.success("✅"); st.rerun()
            if dis: login_prompt_button("form")
            st.markdown("---")
            st.markdown("**LaTeX Quick Ref**")
            if PD_OK:
                st.dataframe(pd.DataFrame({
                    "Description":["Fraction","Sqrt","Sub/Super","Integral","Sum","Greek"],
                    "LaTeX":[r"\frac{a}{b}",r"\sqrt{x}",r"x_{i}^{2}",
                              r"\int_{a}^{b}f(x)dx",r"\sum_{i=1}^{n}x_i",r"\alpha,\beta,\Delta"],
                }), use_container_width=True, hide_index=True)

        with R:
            st.markdown("**Live Preview**")
            test_l = st.text_input("LaTeX", placeholder=r"E = mc^2", key="ltx_test")
            if test_l:
                try: st.latex(test_l)
                except Exception as e: st.error(str(e))
            st.markdown("---")
            st.markdown(f'<div class="sec-hd">📋 {t("w_forms")} ({len(st.session_state.formulas)})</div>',
                        unsafe_allow_html=True)
            if not st.session_state.formulas:
                st.info(t("no_forms"))
            else:
                for i, frm in enumerate(st.session_state.formulas):
                    lbl = frm["desc"][:40] if frm["desc"] else frm["latex"][:30]
                    with st.expander(f"🧮 ({frm['number']}) {lbl}"):
                        try: st.latex(frm["latex"])
                        except: st.code(frm["latex"])
                        if frm.get("desc"): st.caption(frm["desc"])
                        if not dis and st.button(f"🗑️ {t('del_btn')}", key=f"dfm{i}"):
                            st.session_state.formulas.pop(i); st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
#  BIBTEX + CITATION FORMATTERS
# ═══════════════════════════════════════════════════════════════════════════
def _fmt_au(raw):
    parts = [a.strip() for a in re.split(r'\s+and\s+',raw,flags=re.IGNORECASE)]
    out   = []
    for p in parts:
        if not p: continue
        if "," in p:
            segs=p.split(",",1); last=segs[0].strip(); first=segs[1].strip() if len(segs)>1 else ""
            inits="".join(w[0].upper()+"." for w in first.split() if w)
            out.append(f"{last} {inits}".strip())
        else:
            words=p.split()
            if len(words)>=2:
                out.append(f"{words[-1]} "+"".join(w[0].upper()+"." for w in words[:-1]))
            else: out.append(p)
    return ", ".join(out)

def parse_bibtex(text):
    TM={"article":"Journal Article","book":"Book","inbook":"Book Chapter",
        "incollection":"Book Chapter","inproceedings":"Conference Paper",
        "conference":"Conference Paper","misc":"Website",
        "phdthesis":"Thesis","mastersthesis":"Thesis"}
    results=[]
    for raw in re.split(r'(?=@\w+\s*[\{\(])',text.strip()):
        raw=raw.strip()
        if not raw or not raw.startswith("@"): continue
        tm=re.match(r'@(\w+)\s*[\{\(]',raw,re.IGNORECASE)
        if not tm: continue
        et=tm.group(1).lower()
        if et in ("comment","string","preamble"): continue
        fields={}
        for m in re.finditer(r'(\w+)\s*=\s*\{((?:[^{}]|\{[^{}]*\})*)\}',raw,re.DOTALL):
            k2=m.group(1).lower()
            v2=re.sub(r'\s+',' ',m.group(2).strip()).replace("{","").replace("}","")
            fields[k2]=v2
        for m in re.finditer(r'(\w+)\s*=\s*"([^"]*)"',raw,re.DOTALL):
            k2=m.group(1).lower()
            if k2 not in fields: fields[k2]=m.group(2).strip()
        for m in re.finditer(r'(\w+)\s*=\s*(\d{4})\b',raw):
            k2=m.group(1).lower()
            if k2 not in fields: fields[k2]=m.group(2)
        title=fields.get("title","").strip()
        if not title: continue
        results.append({"type":TM.get(et,"Journal Article"),"authors":_fmt_au(fields.get("author","")),
            "year":fields.get("year",""),"title":title,
            "journal":fields.get("journal",fields.get("booktitle",fields.get("series",""))),
            "volume":fields.get("volume",""),"number":fields.get("number",""),
            "pages":fields.get("pages","").replace("--","–"),
            "doi":fields.get("doi",""),"city":fields.get("address",""),
            "publisher":fields.get("publisher","")})
    return results

def fmt_ref(ref, style, n):
    au=ref.get("authors",""); yr=ref.get("year",""); ti=ref.get("title","")
    jn=ref.get("journal",""); vo=ref.get("volume",""); no=ref.get("number","")
    pp=ref.get("pages",""); doi=ref.get("doi","")
    ds=f" DOI: {doi}" if doi else ""
    if "APA" in style:
        vn=f"*{vo}*"+(f"({no})" if no else "") if vo else ""
        return f"{au} ({yr}). {ti}. *{jn}*, {vn}, {pp}.{ds}"
    if "Vancouver" in style or "Ванкувер" in style:
        return f"{n}. {au}. {ti}. {jn}. {yr};{vo}({no}):{pp}.{ds}"
    if "IEEE" in style:
        return f'[{n}] {au}, "{ti}," *{jn}*, vol. {vo}, no. {no}, pp. {pp}, {yr}.{ds}'
    return f"{au} ({yr}). {ti}. *{jn}*, {vo}({no}), {pp}.{ds}"

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: REFS
# ═══════════════════════════════════════════════════════════════════════════
def pg_refs():
    st.markdown(f"## 📑 {t('nav_refs')}")
    guest_banner(); inline_nav(); st.write("")
    dis = not is_logged_in()
    cs_list = t("cite_styles")
    L, R = st.columns(2, gap="large")

    with L:
        st.markdown(f'<div class="sec-hd">➕ {t("add_ref")}</div>', unsafe_allow_html=True)
        rt_list = t("ref_types")
        with st.form("rff", clear_on_submit=True):
            rtype = st.selectbox(t("ref_type"), rt_list, disabled=dis)
            rau   = st.text_input(t("ref_au"),  placeholder="Samarkhanov K.B., Doe J.", disabled=dis)
            ryr   = st.text_input(t("ref_yr"),  placeholder="2024", disabled=dis)
            rti   = st.text_input(t("ref_ti"),  placeholder="Article title", disabled=dis)
            rjn   = st.text_input(t("ref_jn"),  placeholder="Remote Sensing", disabled=dis)
            c1,c2,c3 = st.columns(3)
            rvo = c1.text_input(t("ref_vol"),placeholder="15",disabled=dis)
            rno = c2.text_input(t("ref_no"), placeholder="3", disabled=dis)
            rpp = c3.text_input(t("ref_pp"), placeholder="1234–1250",disabled=dis)
            rdoi = st.text_input(t("ref_doi"), placeholder="10.3390/...",disabled=dis)
            rcit = rpub = ""
            if rtype in [rt_list[1], rt_list[-1]]:   # Book or Thesis
                cx1,cx2 = st.columns(2)
                rcit=cx1.text_input(t("ref_city"),disabled=dis)
                rpub=cx2.text_input(t("ref_pub"), disabled=dis)
            if st.form_submit_button(f"➕ {t('add_btn')}",use_container_width=True,disabled=dis) and rti:
                st.session_state.refs.append({
                    "type":rtype,"authors":rau,"year":ryr,"title":rti,"journal":rjn,
                    "volume":rvo,"number":rno,"pages":rpp,"doi":rdoi,"city":rcit,"publisher":rpub})
                st.success("✅")
        if dis: login_prompt_button("ref")
        st.divider()
        st.markdown(f'<div class="sec-hd">📥 {t("import_bibtex")}</div>', unsafe_allow_html=True)
        bib = st.text_area("BibTeX",height=140,
                           placeholder="@article{key2024,\n  author={Doe, J.},\n  ...}",
                           key="bib_in",disabled=dis)
        ca,cb = st.columns(2)
        if ca.button(t("import_bibtex"),use_container_width=True,disabled=dis):
            if bib.strip():
                p = parse_bibtex(bib)
                if p: st.session_state.refs.extend(p); st.success(f"✅ {len(p)} {t('ref_s')}")
                else: st.warning("No valid entries.")
        if cb.button(t("import_plain"),use_container_width=True,disabled=dis):
            lines=[ln.strip() for ln in bib.split("\n") if ln.strip()]
            for ln in lines:
                st.session_state.refs.append({"type":rt_list[0],"authors":"","year":"",
                    "title":ln,"journal":"","volume":"","number":"","pages":"","doi":"","city":"","publisher":""})
            if lines: st.success(f"✅ {len(lines)} {t('ref_s')}")

    with R:
        style = st.selectbox(t("cite_style_lbl"), cs_list,
                             index=cs_list.index(st.session_state.cite_style)
                             if st.session_state.cite_style in cs_list else 0,
                             key="rs_sel")
        st.session_state.cite_style = style
        st.markdown(f'<div class="sec-hd">📋 {t("w_refs")} ({len(st.session_state.refs)})</div>',
                    unsafe_allow_html=True)
        if not st.session_state.refs: st.info(t("no_refs"))
        else:
            for i,ref in enumerate(st.session_state.refs):
                c1,c2=st.columns([11,1])
                c1.markdown(f'<div class="ri"><span class="rn">[{i+1}]</span> '
                            f'{fmt_ref(ref,style,i+1)}</div>',unsafe_allow_html=True)
                if not dis and c2.button("🗑️",key=f"dr{i}"):
                    st.session_state.refs.pop(i); st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
#  DOCX BUILDER
# ═══════════════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════════════
#  ENU TEMPLATE PATHS  &  MDPI STYLE HELPERS
# ═══════════════════════════════════════════════════════════════════════════
_TPL_PATHS = {
    "EN": "/mnt/user-data/uploads/English_template_2025__1_-OTH.docx",
    "RU": "/mnt/user-data/uploads/Russian_template_2025__1_-OTH.docx",
    "KZ": "/mnt/user-data/uploads/Kazakh_template_2025__1_-OTH.docx",
}

def _tpl_lang() -> str:
    return current_template().get("lang", "EN")

def _safe_clear(doc):
    """Clear body content but preserve sectPr so sections/margins work."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    body = doc.element.body
    children = list(body)
    sectPr = children[-1] if children and children[-1].tag == qn('w:sectPr') else None
    for el in list(body):
        body.remove(el)
    if sectPr is not None:
        body.append(sectPr)
    else:
        body.append(OxmlElement('w:sectPr'))
    return doc

def _add_table_grid(doc, rows: int, cols: int):
    """Add a table with simple single-line borders (no named style needed)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=rows, cols=cols)
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    return table

# ── Multilingual labels ───────────────────────────────────────────────────
def _lbl(key: str) -> str:
    lk = _lang_key()
    _L = {
        "abstract":         {"kz":"Аңдатпа",             "ru":"Аннотация",            "en":"Abstract"},
        "keywords":         {"kz":"Түйін сөздер",         "ru":"Ключевые слова",        "en":"Keywords"},
        "figure":           {"kz":"Сурет",                "ru":"Рисунок",              "en":"Figure"},
        "table":            {"kz":"Кесте",                "ru":"Таблица",              "en":"Table"},
        "formula":          {"kz":"Формула",              "ru":"Формула",              "en":"Formula"},
        "references":       {"kz":"Әдебиеттер тізімі",   "ru":"Список литературы",    "en":"References"},
        "supp":             {"kz":"Қосымша материалдар",  "ru":"Вспомогательный материал","en":"Supplementary Materials"},
        "author_contrib":   {"kz":"Авторлардың үлесі",   "ru":"Вклад авторов",        "en":"Author Contributions"},
        "author_info":      {"kz":"Авторлар туралы",      "ru":"Об авторах",           "en":"Author Information"},
        "funding":          {"kz":"Қаржыландыру",         "ru":"Финансирование",       "en":"Funding"},
        "acknowledgements": {"kz":"Алғыс",               "ru":"Благодарности",        "en":"Acknowledgements"},
        "conflicts":        {"kz":"Мүдделер қақтығысы",  "ru":"Конфликт интересов",   "en":"Conflicts of Interest"},
        "no_funding":       {"kz":"Бұл зерттеу сыртқы қаржыландыруды алмады.",
                             "ru":"Данное исследование не получало внешнего финансирования.",
                             "en":"This research received no external funding."},
        "no_conflicts":     {"kz":"Авторлар мүдделер қақтығысы жоқ деп мәлімдейді.",
                             "ru":"Авторы заявляют об отсутствии конфликта интересов.",
                             "en":"The authors declare no conflicts of interest."},
    }
    return _L.get(key, {}).get(lk, _L.get(key, {}).get("en", key))

def _fig_words(lk: str) -> list:
    return {"kz":["Сурет","сурет"],"ru":["Рисунок","рисунок"],"en":["Figure","figure"]}[lk]

def _tbl_words(lk: str) -> list:
    return {"kz":["Кесте","кесте"],"ru":["Таблица","таблица"],"en":["Table","table"]}[lk]

# ── MDPI style paragraph adders ───────────────────────────────────────────
def _mdpi(doc, style: str, text: str = "", bold_prefix: str = ""):
    """Add paragraph with named MDPI style. bold_prefix = bold run before text."""
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    if text:
        p.add_run(text)
    return p

def _mdpi_heading(doc, number: int, label: str, style="MDPI_2.1_heading1"):
    lbl = f"{number}. {label}" if number else label
    p = doc.add_paragraph(lbl, style=style)
    return p

def _mdpi_backmatter(doc, number: int, label: str, content: str):
    p = doc.add_paragraph(style="MDPI_6.2_BackMatter")
    r = p.add_run(f"{number}. {label}: ")
    r.bold = True
    if content:
        p.add_run(content)
    return p

# ── Figure insertion ──────────────────────────────────────────────────────
def _insert_figure(doc, fig: dict, lk: str):
    fig_word = _lbl("figure")
    if fig.get("image"):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(BytesIO(fig["image"]), width=Inches(4.5))
        except Exception:
            pass
    cap_p = doc.add_paragraph(style="MDPI_5.1_figure_caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl = cap_p.add_run(f"{fig_word} {fig['number']}. ")
    r_lbl.bold = True
    r_cap = cap_p.add_run(fig.get("caption", ""))
    r_cap.italic = True

# ── Table insertion ───────────────────────────────────────────────────────
def _insert_table(doc, tbl: dict, lk: str):
    tbl_word = _lbl("table")
    cap_p = doc.add_paragraph(style="MDPI_4.1_table_caption")
    r_lbl = cap_p.add_run(f"{tbl_word} {tbl['number']}. ")
    r_lbl.bold = True
    cap_p.add_run(tbl.get("caption", ""))
    if tbl.get("data") and PD_OK:
        try:
            df = pd.read_csv(io.StringIO(tbl["data"]))
            wt = _add_table_grid(doc, len(df)+1, len(df.columns))
            for ci, col in enumerate(df.columns):
                c = wt.cell(0, ci)
                c.text = str(col)
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].bold = True
            for ri in range(len(df)):
                for ci in range(len(df.columns)):
                    wt.cell(ri+1, ci).text = str(df.iloc[ri, ci])
        except Exception:
            pass
    doc.add_paragraph()

# ── Formula insertion ─────────────────────────────────────────────────────
def _insert_formula(doc, frm: dict, lk: str):
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{frm.get('latex','')}  ({frm['number']})")
    r.font.name = "Cambria Math"
    if frm.get("desc"):
        d = doc.add_paragraph(frm["desc"], style="MDPI_3.2_text_no_indent")
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if d.runs:
            d.runs[0].italic = True

# ── Section body with cross-reference insertion ───────────────────────────
def _add_section_body(doc, raw_text: str, lk: str,
                      ins_figs: set, ins_tbls: set, ins_forms: set):
    """
    Write section paragraphs. After each paragraph that contains a citation marker,
    insert the cited object inline. Cross-reference markers:
      Figure N / Рисунок N / Сурет N  → inserts the figure
      Table N / Таблица N / Кесте N   → inserts the table
      (N)                               → inserts formula N
    Uncited objects are appended at end of section.
    """
    figs  = {str(f["number"]): f for f in st.session_state.figures}
    tbls  = {str(t_["number"]): t_ for t_ in st.session_state.tables}
    forms = {str(f["number"]): f for f in st.session_state.formulas}
    fw = _fig_words(lk)
    tw = _tbl_words(lk)

    for para_text in raw_text.split("\n"):
        if not para_text.strip():
            doc.add_paragraph(style="MDPI_3.1_text")
            continue
        p = doc.add_paragraph(para_text, style="MDPI_3.1_text")

        # After paragraph: insert cited figures
        for w in fw:
            for m in re.finditer(rf'\b{re.escape(w)}\s+(\d+)\b', para_text):
                n = m.group(1)
                if n in figs and n not in ins_figs:
                    _insert_figure(doc, figs[n], lk)
                    ins_figs.add(n)
        # Tables
        for w in tw:
            for m in re.finditer(rf'\b{re.escape(w)}\s+(\d+)\b', para_text):
                n = m.group(1)
                if n in tbls and n not in ins_tbls:
                    _insert_table(doc, tbls[n], lk)
                    ins_tbls.add(n)
        # Formulas: (N) pattern
        for m in re.finditer(r'\((\d+)\)', para_text):
            n = m.group(1)
            if n in forms and n not in ins_forms:
                _insert_formula(doc, forms[n], lk)
                ins_forms.add(n)

    # Append any not-yet-cited objects at end of section
    for n, fig in figs.items():
        if n not in ins_figs:
            _insert_figure(doc, fig, lk)
            ins_figs.add(n)
    for n, tbl in tbls.items():
        if n not in ins_tbls:
            _insert_table(doc, tbl, lk)
            ins_tbls.add(n)
    for n, frm in forms.items():
        if n not in ins_forms:
            _insert_formula(doc, frm, lk)
            ins_forms.add(n)

# ═══════════════════════════════════════════════════════════════════════════
#  BUILD DOCX — Full ENU/MDPI Template Compliance
# ═══════════════════════════════════════════════════════════════════════════
def build_docx() -> BytesIO:
    """
    Generate ENU-journal-compliant DOCX using the actual ENU template as base.
    Structure follows ENU/MDPI standard:
      IRSTI · Section · Article Type
      Title (MDPI_1.2_title)
      Authors (MDPI_1.3_authornames)
      Affiliations (MDPI_1.6_affiliation)
      Abstract (MDPI_1.7_abstract)  ← bold label
      Keywords (MDPI_1.8_keywords)  ← bold label
      1. Intro … 5. Conclusion  (MDPI_2.1_heading1 + MDPI_3.1_text)
        └─ figures/tables/formulas inserted inline at first citation
      6. Supplementary … 11. Conflicts  (MDPI_6.2_BackMatter)
      12. References  (MDPI_7.1_References, hanging indent, 11pt)
      [EN only] KZ + RU trilingual abstracts
    """
    lk     = _lang_key()
    tpl_lg = _tpl_lang()
    tpl    = current_template()

    # ── Load matching ENU template (inherits all MDPI styles + fonts) ─────
    tpl_path = _TPL_PATHS.get(tpl_lg, _TPL_PATHS["EN"])
    try:
        doc = Document(tpl_path)
        _safe_clear(doc)
    except Exception:
        doc = Document()

    # ── ENU standard margins: 2 cm all sides ─────────────────────────────
    from docx.util import Cm
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2)
        sec.right_margin  = Cm(2)

    # ════ 1. HEADER BLOCK ════════════════════════════════════════════════
    irsti       = st.session_state.get("irsti", "")
    irsti_label = tpl.get("irsti_label", "IRSTI")
    sect_field  = st.session_state.get("section_field", "")
    art_type    = st.session_state.get("art_type", "")

    if irsti_label and irsti:
        doc.add_paragraph(f"{irsti_label} {irsti}")
    if sect_field:
        lbl_sec = {"kz":"Секция","ru":"Секция","en":"Section"}[lk]
        doc.add_paragraph(f"{lbl_sec}: {sect_field}")
    if art_type:
        doc.add_paragraph(art_type, style="MDPI_1.1_article_type")

    # ════ 2. TITLE ═══════════════════════════════════════════════════════
    doc.add_paragraph(st.session_state.art_title or "Untitled",
                      style="MDPI_1.2_title")

    # ════ 3. AUTHORS ══════════════════════════════════════════════════════
    if st.session_state.authors:
        doc.add_paragraph(st.session_state.authors, style="MDPI_1.3_authornames")

    # ════ 4. AFFILIATIONS ═════════════════════════════════════════════════
    if st.session_state.affiliation:
        for line in st.session_state.affiliation.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="MDPI_1.6_affiliation")

    # ════ 5. ABSTRACT ═════════════════════════════════════════════════════
    if st.session_state.abstract:
        abs_lbl = _lbl("abstract")
        p_abs = doc.add_paragraph(style="MDPI_1.7_abstract")
        r_lbl = p_abs.add_run(f"{abs_lbl}: ")
        r_lbl.bold = True
        p_abs.add_run(st.session_state.abstract)

    # ════ 6. KEYWORDS ══════════════════════════════════════════════════════
    if st.session_state.keywords:
        kw_lbl = _lbl("keywords")
        p_kw = doc.add_paragraph(style="MDPI_1.8_keywords")
        r_kl = p_kw.add_run(f"{kw_lbl}: ")
        r_kl.bold = True
        p_kw.add_run(st.session_state.keywords)

    # ════ 7. MAIN SECTIONS ════════════════════════════════════════════════
    MAIN_SECS = ["intro","materials_methods","results","discussion","conclusion"]
    main_num  = 1

    # Shared cross-reference tracking (figures/tables/formulas insert once)
    _ins_figs  = set()
    _ins_tbls  = set()
    _ins_forms = set()

    for key in active_sections():
        if key not in MAIN_SECS:
            continue
        val = st.session_state.get(key, "").strip()
        if not val:
            main_num += 1
            continue

        _mdpi_heading(doc, main_num, sec_name(key))
        main_num += 1

        _add_section_body(doc, val, lk, _ins_figs, _ins_tbls, _ins_forms)

    # ════ 8. BACK MATTER (numbered, MDPI_6.2_BackMatter style) ═══════════
    back_num = main_num

    def _bm_if(label_key: str, sess_key: str):
        nonlocal back_num
        content = st.session_state.get(sess_key, "").strip()
        if not content: return
        _mdpi_backmatter(doc, back_num, _lbl(label_key), content)
        back_num += 1

    _bm_if("supp",           "supplementary")
    _bm_if("author_contrib", "author_contributions")
    _bm_if("author_info",    "author_info")

    # Funding — mandatory (default text if empty)
    funding_txt = st.session_state.get("funding","").strip() or _lbl("no_funding")
    _mdpi_backmatter(doc, back_num, _lbl("funding"), funding_txt)
    back_num += 1

    _bm_if("acknowledgements", "acknowledgements")

    # Conflicts — mandatory
    conflicts_txt = st.session_state.get("conflicts","").strip() or _lbl("no_conflicts")
    _mdpi_backmatter(doc, back_num, _lbl("conflicts"), conflicts_txt)
    back_num += 1

    # ════ 9. REFERENCES (MDPI_7.1_References, hanging indent) ════════════
    if st.session_state.refs:
        _mdpi_heading(doc, back_num, _lbl("references"))
        back_num += 1
        for i, ref in enumerate(st.session_state.refs, 1):
            p_ref = doc.add_paragraph(style="MDPI_7.1_References")
            # Hanging indent: 0.5 cm
            from docx.util import Cm as _Cm
            p_ref.paragraph_format.left_indent         = _Cm(0.5)
            p_ref.paragraph_format.first_line_indent   = _Cm(-0.5)
            p_ref.paragraph_format.space_after         = Pt(3)
            p_ref.add_run(f"{i}. {fmt_ref(ref, st.session_state.cite_style, i)}")

    # ════ 10. EN TEMPLATE: TRILINGUAL ABSTRACTS (KZ + RU) ════════════════
    if tpl_lg == "EN":
        for tri_lang, abs_key, abs_lbl, kw_lbl in [
            ("KZ", "abstract_kz", "Аңдатпа",   "Түйін сөздер"),
            ("RU", "abstract_ru", "Аннотация", "Ключевые слова"),
        ]:
            tri_abs = st.session_state.get(abs_key, "").strip()
            if not tri_abs:
                continue
            doc.add_paragraph()
            # Separator
            sep = doc.add_paragraph("─" * 50)
            sep.paragraph_format.space_before = Pt(12)

            doc.add_paragraph(st.session_state.art_title or "",
                              style="MDPI_1.2_title")
            if st.session_state.authors:
                doc.add_paragraph(st.session_state.authors,
                                  style="MDPI_1.3_authornames")
            p_abs = doc.add_paragraph(style="MDPI_1.7_abstract")
            r = p_abs.add_run(f"{abs_lbl}: ")
            r.bold = True
            p_abs.add_run(tri_abs)
            if st.session_state.keywords:
                p_kw = doc.add_paragraph(style="MDPI_1.8_keywords")
                r2 = p_kw.add_run(f"{kw_lbl}: ")
                r2.bold = True
                p_kw.add_run(st.session_state.keywords)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: GENERATE / EXPORT
# ═══════════════════════════════════════════════════════════════════════════
def pg_generate():
    st.markdown(f"## 🚀 {t('nav_gen')}")
    guest_banner(); inline_nav(); st.write("")
    dis=not is_logged_in()
    score,missing=article_readiness()

    st.markdown(f"### 📋 {t('readiness')}")
    cp,cs,cr,cf=st.columns(4)
    cp.progress(score/100, text=f"{score}%")
    cs.metric(t("w_secs"),f"{sum(1 for k in active_sections() if st.session_state.get(k,'').strip())}/{len(active_sections())}")
    cr.metric(t("w_refs"),len(st.session_state.refs))
    cf.metric(t("w_forms"),len(st.session_state.formulas))

    if missing:
        miss_str = ", ".join(sec_name(m) if m in ALL_SECTIONS else m for m in missing)
        st.warning(f"⚠️ {t('missing')}: **{miss_str}**")

    can_export = score >= 70 and is_logged_in()
    if not can_export and is_logged_in():
        st.info("📝 Барлық негізгі бөлімдерді толтырып, дереккөздер қосыңыз — экспорт ашылады.")
    if not is_logged_in():
        login_prompt_button("export")

    # Preview
    st.markdown(f"### {t('preview_lbl')}")
    keys=active_sections()
    c_ = _colors()
    parts=[
        f'<div class="pt">{st.session_state.art_title or "—"}</div>',
        f'<div class="pa"><b>{st.session_state.authors}</b></div>',
        f'<div class="pa" style="font-size:0.78rem;">{st.session_state.affiliation}</div>',
    ]
    if st.session_state.keywords:
        parts.append(f'<div style="text-align:center;font-size:0.75rem;color:{c_["muted"]};"><b>Keywords:</b> {st.session_state.keywords}</div>')
    if st.session_state.abstract:
        ab=st.session_state.abstract
        parts.append(f'<div class="pab"><b>Abstract:</b> {ab[:500]}{"…" if len(ab)>500 else ""}</div>')
    for i,k in enumerate(keys,1):
        c=st.session_state.get(k,"")
        if c: parts.append(f'<div class="ph">{i}. {sec_name(k)}</div><p style="font-size:0.84rem;">{c[:400]}{"…" if len(c)>400 else ""}</p>')
    if st.session_state.refs:
        parts.append('<div class="ph">References</div>')
        for j,ref in enumerate(st.session_state.refs,1):
            parts.append(f'<p style="font-size:0.79rem;">{fmt_ref(ref,st.session_state.cite_style,j)}</p>')
    st.markdown(f'<div class="pv">{"".join(parts)}</div>',unsafe_allow_html=True)
    st.write("")

    st.markdown(f"### 📥 Export")
    fn=sfn(st.session_state.art_title)
    e1,e2,e3=st.columns(3)

    with e1:
        if DOCX_OK and can_export:
            buf=build_docx()
            st.download_button(t("dl_docx"),buf,f"{fn}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        else:
            st.button(t("dl_docx"),disabled=True,use_container_width=True,
                      help="Барлық бөлімдерді толтырыңыз" if is_logged_in() else t("login_required"))

    with e2:
        md=["# "+st.session_state.art_title,
            f"**{st.session_state.authors}**",f"*{st.session_state.affiliation}*",
            "",f"**Keywords:** {st.session_state.keywords}","",
            f"## Abstract\n\n{st.session_state.abstract}",""]
        for i,k in enumerate(active_sections(),1):
            if st.session_state.get(k,""): md+=[f"## {i}. {sec_name(k)}",st.session_state[k],""]
        if can_export:
            st.download_button(t("dl_md"),"\n".join(md).encode("utf-8"),
                f"{fn}.md","text/markdown",use_container_width=True)
        else:
            st.button(t("dl_md"),disabled=True,use_container_width=True)

    with e3:
        proj={k:st.session_state.get(k,"") for k in
              list(ALL_SECTIONS)+list(THESIS_SECTIONS)+
              ["art_title","authors","affiliation","journal","keywords","abstract",
               "art_type","irsti","section_field","template","cite_style"]}
        proj["refs"]=st.session_state.refs
        proj["formulas"]=st.session_state.formulas
        proj["tables"]=st.session_state.tables
        proj["exported_at"]=datetime.now().isoformat()
        if is_logged_in():
            st.download_button(t("save_json"),
                json.dumps(proj,ensure_ascii=False,indent=2).encode("utf-8"),
                f"{fn}_project.json","application/json",use_container_width=True)
        else:
            st.button(t("save_json"),disabled=True,use_container_width=True)

    st.markdown("---")
    st.subheader(t("feedback_lbl"))
    with st.form("fb",clear_on_submit=True):
        fb=st.text_area("",height=80,placeholder=t("feedback_ph"),
                        label_visibility="collapsed",disabled=dis)
        if st.form_submit_button(t("feedback_send"),use_container_width=True,disabled=dis):
            if fb.strip():
                add_log("feedback",st.session_state.username,fb[:300])
                send_email_notification(
                    f"Feedback — Smart Article: {st.session_state.username}",
                    f"<h2>New Feedback</h2><p><b>From:</b> {st.session_state.username}</p>"
                    f"<p><b>Message:</b><br>{fb}</p>"
                    f"<p><b>Time:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>")
                st.success(t("feedback_ok"))

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: THESIS / DIPLOMA
# ═══════════════════════════════════════════════════════════════════════════
def pg_thesis():
    lk=_lang_key()
    title_map={"kz":"🎓 Дипломдық жұмыс","ru":"🎓 Дипломная работа","en":"🎓 Diploma / Thesis"}
    st.markdown(f"## {title_map[lk]}")
    guest_banner(); inline_nav(); st.write("")
    dis=not is_logged_in()

    T_META, T_MAIN, T_EXTRA, T_CHECK = st.tabs([
        {"kz":"📋 Мәліметтер","ru":"📋 Данные","en":"📋 Metadata"}[lk],
        {"kz":"✍️ Тараулар","ru":"✍️ Главы","en":"✍️ Chapters"}[lk],
        {"kz":"📎 Қосымшалар","ru":"📎 Приложения","en":"📎 Appendices"}[lk],
        {"kz":"📏 Нормоконтроль","ru":"📏 Нормоконтроль","en":"📏 Normative Check"}[lk],
    ])

    with T_META:
        st.markdown(f'<div class="sec-hd">{"Дипломдық жұмыс туралы мәліметтер" if lk=="kz" else "Данные дипломной работы" if lk=="ru" else "Diploma Work Metadata"}</div>',unsafe_allow_html=True)
        col1,col2=st.columns(2)
        for key,lbl,ph in [
            ("th_student",{"kz":"Студент","ru":"Студент","en":"Student"}[lk],"Lastname Firstname"),
            ("th_supervisor",{"kz":"Ғылыми жетекші","ru":"Научный руководитель","en":"Supervisor"}[lk],"Prof. Lastname F."),
            ("th_university",{"kz":"Университет","ru":"Университет","en":"University"}[lk],"L.N. Gumilyov ENU"),
            ("th_faculty",{"kz":"Факультет","ru":"Факультет","en":"Faculty"}[lk],"Faculty of Natural Sciences"),
            ("th_specialty",{"kz":"Мамандық","ru":"Специальность","en":"Speciality"}[lk],"6B05301 — Geography"),
            ("th_degree",{"kz":"Дәреже","ru":"Степень","en":"Degree"}[lk],"Bachelor / Master / PhD"),
            ("th_year",{"kz":"Жыл","ru":"Год","en":"Year"}[lk],"2025"),
        ]:
            col=(col1 if list(["th_student","th_university","th_specialty","th_year"]).count(key) else col2)
            v=col.text_input(lbl,value=st.session_state.get(key,""),placeholder=ph,disabled=dis)
            if not dis: st.session_state[key]=v

        if dis: login_prompt_button("thesis_meta")
        st.divider()
        # Title page preview
        st.markdown(f'<div class="pv" style="padding:16px;text-align:center;">'
            f'<div style="font-size:0.8rem;margin-bottom:8px;">{st.session_state.th_university}</div>'
            f'<div style="font-size:0.8rem;margin-bottom:16px;">{st.session_state.th_faculty}</div>'
            f'<div style="font-weight:700;font-size:1.1rem;margin-bottom:8px;">{st.session_state.art_title or "Дипломдық жұмыстың тақырыбы"}</div>'
            f'<div style="font-size:0.8rem;margin-bottom:4px;">{st.session_state.th_specialty}</div>'
            f'<div style="font-size:0.8rem;margin-bottom:16px;">{st.session_state.th_degree}</div>'
            f'<div style="font-size:0.8rem;">{"Орындаған" if lk=="kz" else "Выполнил" if lk=="ru" else "Prepared by"}: '
            f'{st.session_state.th_student}</div>'
            f'<div style="font-size:0.8rem;">{"Ғылыми жетекші" if lk=="kz" else "Научный руководитель" if lk=="ru" else "Supervisor"}: '
            f'{st.session_state.th_supervisor}</div>'
            f'<div style="font-size:0.8rem;margin-top:16px;">{st.session_state.th_year}</div>'
            f'</div>', unsafe_allow_html=True)

    with T_MAIN:
        th_main_keys = ["th_abstract_kz","th_abstract_ru","th_abstract_en","th_abbrev",
                        "th_intro","th_ch1","th_ch2","th_ch3","th_ch4","th_concl"]
        tab_labels=[f"{THESIS_SECTIONS.get(k,{}).get('icon','•')} {sec_name(k)}" for k in th_main_keys]
        th_tabs=st.tabs(tab_labels)
        for tab,key in zip(th_tabs,th_main_keys):
            with tab:
                st.markdown(f'<div class="sec-hd">{sec_name(key)}</div>',unsafe_allow_html=True)
                # Hints per chapter
                hints_map={
                    "th_intro":{"kz":"Кіріспеде: тақырыптың өзектілігі, зерттеу мақсаты, міндеттер, жаңалық, практикалық мәні, зерттеу объектісі мен пәні.","ru":"Введение: актуальность темы, цель, задачи, научная новизна, практическая значимость, объект и предмет исследования.","en":"Introduction: relevance, objective, tasks, novelty, practical significance, object and subject of research."},
                    "th_ch1":{"kz":"1-тарау: тақырыпқа байланысты негізгі теориялық ережелер мен авторлардың пікірлері. Дереккөздерге сілтемелер міндетті.","ru":"Глава 1: теоретические основы, анализ литературы, сопоставление взглядов авторов. Ссылки обязательны.","en":"Chapter 1: theoretical foundations, literature analysis, comparison of authors' views. References are mandatory."},
                    "th_ch2":{"kz":"2-тарау: зерттеу дизайны, деректер жинау әдістемесі, талдау жасалатын деректер базасы.","ru":"Глава 2: методология исследования, методы сбора и обработки данных, описание базы данных.","en":"Chapter 2: research design, data collection methodology, database description."},
                    "th_ch3":{"kz":"3-тарау: нәтижелер, кестелер мен суреттерге сілтемелер, алынған деректерді интерпретациялау.","ru":"Глава 3: результаты анализа, ссылки на таблицы и рисунки, интерпретация данных.","en":"Chapter 3: analysis results, references to tables and figures, data interpretation."},
                    "th_concl":{"kz":"Қорытынды: мақсаттың орындалғаны, негізгі нәтижелер, ұсыныстар, болашақ зерттеулер.","ru":"Заключение: достижение цели, основные результаты, рекомендации, перспективы.","en":"Conclusion: achievement of objective, key results, recommendations, future research."},
                }
                hint_txt=hints_map.get(key,{}).get(lk,"")
                if hint_txt:
                    st.markdown(f'<div class="hint-box"><div class="hint-lbl">✦ Smart Article</div>{hint_txt}</div>',unsafe_allow_html=True)

                insert_buttons(key)
                val=st.text_area(sec_name(key),value=st.session_state.get(key,""),
                                 height=320,key=f"th_ta_{key}",disabled=dis,
                                 label_visibility="collapsed",
                                 placeholder=f"{sec_name(key)}…")
                if not dis: st.session_state[key]=val
                st.caption(f"📝 {t('word_count')}: {wc(st.session_state.get(key,''))}")
                if dis: login_prompt_button(f"thesis_{key}")

                # Objectives checker in thesis intro
                if key=="th_intro" and is_logged_in():
                    tmp=st.session_state.get("th_intro","")
                    st.session_state["intro"]=tmp  # temp alias for checker
                    st.markdown("---")
                    st.markdown(f"**{t('objectives_tool')}**")
                    objectives_checker()

    with T_EXTRA:
        L,R=st.columns(2,gap="large")
        for col,key in zip([L,R],["th_appendix_a","th_appendix_b"]):
            with col:
                st.markdown(f'<div class="sec-hd">{sec_name(key)}</div>',unsafe_allow_html=True)
                v=st.text_area(sec_name(key),value=st.session_state.get(key,""),
                               height=200,key=f"th_ta_{key}",disabled=dis,
                               label_visibility="collapsed",
                               placeholder="Appendix content…")
                if not dis: st.session_state[key]=v
                uf=st.file_uploader("TXT/DOCX",type=["txt","docx"],key=f"th_up_{key}",disabled=dis,
                                    label_visibility="collapsed")
                if uf and not dis:
                    if uf.name.endswith(".txt"): st.session_state[key]=uf.read().decode("utf-8",errors="ignore")
                    elif uf.name.endswith(".docx") and DOCX_OK:
                        try:
                            doc=Document(BytesIO(uf.read()))
                            st.session_state[key]="\n".join(p.text for p in doc.paragraphs)
                        except: pass
        if dis: login_prompt_button("appendix")

    with T_CHECK:
        st.markdown(f'<div class="sec-hd">{t("th_norm_check")}</div>',unsafe_allow_html=True)
        lk2=_lang_key()
        reqs={
            "kz":[
                ("th_pages","Беттер саны","60–100 бет","≥60 бет: дипломдық жұмыстың минималды көлемі"),
                ("th_font","Шрифт","Times New Roman 14pt","Стандартты шрифт"),
                ("th_spacing","Аралық","1.5 интервал","Жол аралығы"),
                ("th_margin","Жиектер","Сол: 30мм, Оң: 15мм, Жоғары/Төменгі: 20мм","ГОСТ жиектері"),
            ],
            "ru":[
                ("th_pages","Объём","60–100 страниц","≥60 стр.: минимальный объём дипломной работы"),
                ("th_font","Шрифт","Times New Roman 14pt","Стандартный шрифт"),
                ("th_spacing","Интервал","Полуторный","Межстрочный интервал"),
                ("th_margin","Поля","Лев: 30мм, Пр: 15мм, В/Н: 20мм","Поля по ГОСТ"),
            ],
            "en":[
                ("th_pages","Pages","60–100 pages","≥60 pages: minimum thesis volume"),
                ("th_font","Font","Times New Roman 14pt","Standard thesis font"),
                ("th_spacing","Spacing","1.5 line spacing","Standard line spacing"),
                ("th_margin","Margins","Left: 30mm, Right: 15mm, Top/Bottom: 20mm","GOST margins"),
            ],
        }
        body_total=sum(wc(st.session_state.get(k,"")) for k in THESIS_SECTIONS)
        est_pages=max(1,round(body_total/250))

        for _,lbl,std,desc in reqs[lk2]:
            ok_check=(est_pages>=60 if "60" in std else True)
            icon="✅" if ok_check else "⚠️"
            css="obj-ok" if ok_check else "obj-miss"
            val_display=f"{est_pages} стр. (приблиз.)" if "page" in lbl.lower() or "бет" in lbl.lower() or "объём" in lbl.lower() else std
            st.markdown(f'<div class="{css}">{icon} {lbl}: <b>{val_display}</b> — {desc}</div>',
                        unsafe_allow_html=True)

        word_total=sum(wc(st.session_state.get(k,"")) for k in list(ALL_SECTIONS)+list(THESIS_SECTIONS))
        st.metric({"kz":"Жалпы сөздер","ru":"Всего слов","en":"Total words"}[lk2], word_total)
        st.metric({"kz":"Бағаланған беттер","ru":"Примерно страниц","en":"Estimated pages"}[lk2], est_pages)

# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: SETTINGS
# ═══════════════════════════════════════════════════════════════════════════
def pg_settings():
    lk=_lang_key()
    st.markdown(f"## ⚙️ {t('nav_set')}")
    inline_nav(); st.write("")
    T1,T2,T3=st.tabs([
        {"kz":"🎨 Тема","ru":"🎨 Тема","en":"🎨 Theme"}[lk],
        {"kz":"📂 Жоба","ru":"📂 Проект","en":"📂 Project"}[lk],
        {"kz":"☁️ GitHub","ru":"☁️ GitHub","en":"☁️ GitHub"}[lk],
    ])

    with T1:
        cs_list=t("cite_styles")
        idx=cs_list.index(st.session_state.cite_style) if st.session_state.cite_style in cs_list else 0
        cs=st.selectbox(t("cite_style_lbl"),cs_list,index=idx)
        st.session_state.cite_style=cs
        if is_logged_in() and st.session_state.username=="admin":
            st.divider()
            st.subheader({"kz":"📋 Белсенділік журналы","ru":"📋 Журнал активности","en":"📋 Activity Log"}[lk])
            logs=[]
            if LOGS_FILE.exists():
                try: logs=json.loads(LOGS_FILE.read_text("utf-8"))
                except: pass
            if logs and PD_OK:
                st.dataframe(pd.DataFrame(logs[-50:][::-1]),use_container_width=True,hide_index=True)

    with T2:
        L2,R2=st.columns(2)
        with L2:
            st.subheader(t("load_json"))
            upf=st.file_uploader("JSON",type="json")
            if upf:
                try:
                    data=json.load(upf)
                    for f_ in list(ALL_SECTIONS)+list(THESIS_SECTIONS)+[
                        "art_title","authors","affiliation","journal","keywords","abstract",
                        "art_type","irsti","section_field","template","refs","tables","formulas","cite_style",
                        "th_student","th_supervisor","th_year","th_university","th_faculty","th_specialty","th_degree"]:
                        if f_ in data: st.session_state[f_]=data[f_]
                    st.success(t("loaded"))
                except Exception as e: st.error(str(e))
        with R2:
            st.subheader(t("reset_btn"))
            st.warning({"kz":"Сессиядағы барлық мақала деректері жойылады.","ru":"Все данные статьи в текущей сессии будут удалены.","en":"All article data in current session will be erased."}[lk])
            if st.button(t("reset_btn"),type="secondary"):
                for k in list(ALL_SECTIONS)+list(THESIS_SECTIONS)+[
                    "art_title","authors","affiliation","journal","keywords","abstract",
                    "art_type","irsti","section_field"]:
                    st.session_state[k]=""
                st.session_state.figures=[]; st.session_state.tables=[]
                st.session_state.refs=[]; st.session_state.formulas=[]
                st.success(t("reset_ok"))

    with T3:
        st.subheader(t("gh_title"))
        if not REQ_OK:
            st.error("pip install requests")
        else:
            c1,c2=st.columns(2)
            gh_tok=c1.text_input(t("gh_token"),value=st.session_state.gh_token,
                                  type="password",placeholder="ghp_xxxxxxxxxxxx")
            gh_rep=c2.text_input(t("gh_repo"),value=st.session_state.gh_repo,
                                  placeholder="username/repo")
            st.session_state.gh_token=gh_tok
            st.session_state.gh_repo=gh_rep

            if st.button({"kz":"🧪 Байланысты тексеру","ru":"🧪 Проверить подключение","en":"🧪 Test connection"}[lk]):
                if gh_tok and gh_rep:
                    try:
                        r=requests.get(f"https://api.github.com/repos/{gh_rep}",
                                        headers={"Authorization":f"token {gh_tok}"},timeout=8)
                        if r.status_code==200:
                            st.success(f"✅ Connected: {r.json().get('full_name')}")
                        else:
                            st.error(f"❌ {r.status_code}: {r.json().get('message','')}")
                    except Exception as e: st.error(str(e))
                else:
                    st.warning({"kz":"Token мен репозиторий атын енгізіңіз.","ru":"Введите token и имя репозитория.","en":"Enter token and repository name."}[lk])

            st.markdown("---")
            st.caption({"kz":"Жазылу/кіру тіркелімдері мен логдары GitHub репозиторийіне автоматты түрде синхрондалады. 'data/users.json' мен 'data/logs.json' файлдары жасалады.","ru":"Данные регистраций и логи сессий автоматически синхронизируются с GitHub. Создаются файлы 'data/users.json' и 'data/logs.json'.","en":"Registration data and session logs are automatically synced to GitHub. Files 'data/users.json' and 'data/logs.json' will be created."}[lk])

# ═══════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════════════
#  PAGE: AI AGENT  (full article analysis)
# ═══════════════════════════════════════════════════════════════════════════
def pg_agent():
    lk  = _lang_key()
    st.markdown(f"## 🤖 {t('nav_agent')}")
    guest_banner(); inline_nav(); st.write("")

    if not is_logged_in():
        login_prompt_button("agent"); return

    ai_key = st.session_state.get("anthropic_api_key", "")
    c_ = _colors()

    # ── API key prompt ────────────────────────────────────────────
    if not ai_key:
        st.markdown(
            f'<div class="ai-badge">🤖 '
            f'<a href="https://console.anthropic.com/settings/keys" target="_blank">'
            f'{"✦ ЖИ белсендіру — API кілтін сайдбарға енгізіңіз" if lk=="kz" else "✦ Активировать ИИ — введите API ключ в боковой панели" if lk=="ru" else "✦ Enable AI — enter API key in the sidebar"}'
            f'</a></div>', unsafe_allow_html=True)
        return

    # ── Agent description ─────────────────────────────────────────
    desc = {
        "kz": ("🤖 **ЖИ Агент** — барлық жазылған бөлімдерді біріктіріп, "
               "ғылыми логиканы, стильді және толықтықты жан-жақты талдайды. "
               "Агент нақты ұсыныстар мен жетіспейтін элементтерді анықтайды."),
        "ru": ("🤖 **ИИ Агент** — объединяет все написанные разделы статьи и "
               "проводит комплексный анализ научной логики, стиля и полноты. "
               "Выдаёт конкретные рекомендации и выявляет недостающие элементы."),
        "en": ("🤖 **AI Agent** — combines all written sections and performs a "
               "comprehensive analysis of scientific logic, style, and completeness. "
               "Provides specific recommendations and identifies missing elements."),
    }[lk]
    st.markdown(desc)
    st.markdown("---")

    # ── Article snapshot ──────────────────────────────────────────
    score, missing = article_readiness()
    has_content = any(st.session_state.get(k,"").strip() for k in active_sections())

    col1, col2, col3 = st.columns(3)
    col1.metric({"kz":"Дайындығы","ru":"Готовность","en":"Readiness"}[lk], f"{score}%")
    col2.metric({"kz":"Бөлімдер","ru":"Разделы","en":"Sections"}[lk],
                f"{sum(1 for k in active_sections() if st.session_state.get(k,'').strip())}/{len(active_sections())}")
    col3.metric({"kz":"Дереккөздер","ru":"Источники","en":"References"}[lk],
                len(st.session_state.refs))

    if not has_content:
        st.warning({"kz":"⚠️ Агентті іске қосу үшін кем дегенде бір бөлімді толтырыңыз.",
                    "ru":"⚠️ Заполните хотя бы один раздел для запуска агента.",
                    "en":"⚠️ Fill at least one section to run the agent."}[lk])
        return

    st.markdown("---")

    # ── Analysis mode selector ────────────────────────────────────
    modes = {
        "kz": ["📋 Толық талдау","🎯 Мақсат & Міндеттер","✍️ Академиялық стиль","🔗 Логикалық байланыс","📚 Дереккөздер жеткіліктілігі"],
        "ru": ["📋 Полный анализ","🎯 Цель & Задачи","✍️ Академический стиль","🔗 Логические связи","📚 Достаточность источников"],
        "en": ["📋 Full Analysis","🎯 Objectives & Tasks","✍️ Academic Style","🔗 Logical Coherence","📚 References Sufficiency"],
    }[lk]
    mode = st.selectbox({"kz":"Талдау режимі","ru":"Режим анализа","en":"Analysis mode"}[lk],
                        modes, key="agent_mode")

    # ── Build context ─────────────────────────────────────────────
    def _build_context() -> str:
        parts = []
        if st.session_state.art_title:
            parts.append(f"TITLE: {st.session_state.art_title}")
        if st.session_state.abstract:
            parts.append(f"ABSTRACT:\n{st.session_state.abstract}")
        for k in active_sections():
            v = st.session_state.get(k, "").strip()
            if v:
                parts.append(f"{sec_name(k).upper()}:\n{v[:3000]}")
        if st.session_state.refs:
            ref_lines = [f"[{i+1}] {r.get('authors','')} ({r.get('year','')}) {r.get('title','')}"
                         for i, r in enumerate(st.session_state.refs[:20])]
            parts.append("REFERENCES:\n" + "\n".join(ref_lines))
        return "\n\n".join(parts)

    # ── System prompts per mode ───────────────────────────────────
    SYS = {
        0: {  # Full
            "kz": "Сіз жоғары білікті ғылыми редакторсыз. Мақаланың барлық бөлімдерін мұқият оқып, жан-жақты талдау жасаңыз:\n1) 💪 Күшті тұстары\n2) ⚠️ Жетіспейтін немесе әлсіз тұстары\n3) 🔗 Бөлімдер арасындағы логикалық байланыс\n4) 📝 Нақты жақсарту ұсыныстары (3-5 маңызды)\nҚазақша жауап беріңіз.",
            "ru": "Вы высококвалифицированный научный редактор. Проведите полный анализ статьи:\n1) 💪 Сильные стороны\n2) ⚠️ Слабые места и пробелы\n3) 🔗 Логическая связность разделов\n4) 📝 Конкретные рекомендации (3-5 важных)\nОтвечайте на русском профессионально.",
            "en": "You are an expert scientific editor. Provide a comprehensive analysis:\n1) 💪 Strengths\n2) ⚠️ Weaknesses and gaps\n3) 🔗 Logical coherence between sections\n4) 📝 Specific actionable recommendations (3-5 key ones)\nBe professional and thorough.",
        },
        1: {  # Objectives
            "kz": "Тек Кіріспе бөліміне назар аударыңыз. Зерттеу мақсаты мен міндеттерінің: 1) нақтылығын 2) өлшенетіндігін 3) бөлімнің қорытындысымен сәйкестігін талдаңыз. Нақты жетіспейтін тұстарды атаңыз. Қазақша жауап беріңіз.",
            "ru": "Сосредоточьтесь только на Введении. Проанализируйте цель и задачи: 1) конкретность 2) измеримость 3) соответствие выводам в Заключении. Укажите конкретные недостатки. Отвечайте на русском.",
            "en": "Focus only on the Introduction. Analyze the objectives and tasks: 1) specificity 2) measurability 3) alignment with Conclusion. Point out specific weaknesses. Reply in English.",
        },
        2: {  # Style
            "kz": "Мақаланың академиялық стилін талдаңыз: 1) Ғылыми терминологияның дұрыстығы 2) Пассив/актив конструкциялар 3) Тіл тазалығы мен анықтығы 4) Академиялық емес тіркестер. Мысалдармен дәлелдеңіз. Қазақша жауап беріңіз.",
            "ru": "Проанализируйте академический стиль статьи: 1) Корректность терминологии 2) Использование пассивных/активных конструкций 3) Ясность и чёткость языка 4) Неакадемические выражения. Приведите примеры из текста. Отвечайте на русском.",
            "en": "Analyze the academic style: 1) Terminology correctness 2) Passive/active voice usage 3) Clarity and precision 4) Non-academic expressions. Provide examples from the text. Reply in English.",
        },
        3: {  # Logic
            "kz": "Мақаланың логикалық байланысын тексеріңіз: 1) Кіріспеден Қорытындыға дейінгі логика 2) Нәтижелер Талқылаумен сәйкестігі 3) Мақсат пен Қорытынды арасындағы байланыс 4) Айтылмаған немесе дәлелденбеген тезистер. Қазақша жауап беріңіз.",
            "ru": "Проверьте логическую связность статьи: 1) Логика от Введения до Заключения 2) Соответствие Результатов и Обсуждения 3) Связь между целью и выводами 4) Недоказанные тезисы. Отвечайте на русском.",
            "en": "Check the logical coherence: 1) Flow from Introduction to Conclusion 2) Alignment of Results and Discussion 3) Connection between objectives and conclusions 4) Unproven claims. Reply in English.",
        },
        4: {  # References
            "kz": "Дереккөздер тізімін талдаңыз: 1) Жеткіліктілігі (тақырыпқа сәйкес) 2) Жылдары (соңғы 5-10 жыл) 3) Сілтемелердің мәтін ішінде пайдаланылуы 4) Жетіспейтін маңызды дереккөздер тақырыптары. Қазақша жауап беріңіз.",
            "ru": "Проанализируйте список источников: 1) Достаточность (соответствие теме) 2) Актуальность (последние 5-10 лет) 3) Использование ссылок в тексте 4) Темы отсутствующих важных источников. Отвечайте на русском.",
            "en": "Analyze the references: 1) Sufficiency (topic relevance) 2) Recency (last 5-10 years) 3) In-text citation usage 4) Topics of missing key references. Reply in English.",
        },
    }

    mode_idx = modes.index(mode)
    sys_prompt = SYS[mode_idx][lk]

    # ── Run button ────────────────────────────────────────────────
    run_lbl = {
        "kz": f"🚀 Агентті іске қосу — {mode}",
        "ru": f"🚀 Запустить агента — {mode}",
        "en": f"🚀 Run Agent — {mode}",
    }[lk]

    result_key = f"agent_result_{mode_idx}"

    if st.button(run_lbl, use_container_width=True, type="primary"):
        context = _build_context()
        if not context.strip():
            st.warning("No content to analyze.")
        else:
            spinner_msg = {"kz":"🤖 Агент талдап жатыр...","ru":"🤖 Агент анализирует...","en":"🤖 Agent is analyzing..."}[lk]
            with st.spinner(spinner_msg):
                result = _call_anthropic(context[:50000], system=sys_prompt)
            st.session_state[result_key] = result
            add_log("agent_run", st.session_state.username, f"mode={mode_idx}")

    # ── Display result ────────────────────────────────────────────
    cached = st.session_state.get(result_key, "")
    if cached:
        st.markdown("---")
        result_lbl = {"kz":"📊 Нәтиже","ru":"📊 Результат","en":"📊 Result"}[lk]
        st.markdown(f"### {result_lbl}")
        acc_color = _colors()["acc"]
        # Render newlines as <br> inside styled card
        cached_html = cached.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")
        st.markdown(
            f'<div class="pv" style="border-left:4px solid {acc_color};padding:16px 20px;line-height:1.7;">'
            f'{cached_html}'
            f'</div>', unsafe_allow_html=True)

        col_a, col_b = st.columns(2)
        if col_a.button({"kz":"✕ Тазалау","ru":"✕ Очистить","en":"✕ Clear"}[lk],
                        key=f"agent_clear_{mode_idx}"):
            st.session_state[result_key] = ""; st.rerun()

        col_b.download_button(
            {"kz":"📥 Жүктеу","ru":"📥 Скачать","en":"📥 Download"}[lk],
            cached.encode("utf-8"),
            file_name=f"agent_analysis_{mode_idx}.txt",
            mime="text/plain",
            key=f"agent_dl_{mode_idx}",
        )


def main():
    inject_css(dark=st.session_state.dark)
    sidebar()
    routes={
        "info":     pg_info,
        "sections": pg_sections,
        "media":    pg_media,
        "refs":     pg_refs,
        "generate": pg_generate,
        "thesis":   pg_thesis,
        "agent":    pg_agent,
        "settings": pg_settings,
        "auth":     pg_auth,
    }
    routes.get(st.session_state.page, pg_info)()

if __name__=="__main__":
    main()
